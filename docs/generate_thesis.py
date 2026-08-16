# -*- coding: utf-8 -*-
"""
Genereaza documentul Word al lucrarii de disertatie (Capitolele 1-3),
folosind python-docx. Times New Roman 12pt, aliniere justificata,
spatiere 1.5, capitole incepand pe pagina noua, TOC editabil, numerotare
pagini.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

HERE = os.path.dirname(os.path.abspath(__file__))
ARCH_DIAGRAM = os.path.join(HERE, "architecture_diagram.png")
LATENCY_CHART = os.path.join(HERE, "..", "network_sim", "latency_comparison.png")

FONT = "Times New Roman"
CODE_FONT = "Consolas"

doc = Document()

# ---------------------------------------------------------------- pagina ---
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.0)


# ------------------------------------------------------------ utilitare ---
def set_run_font(run, name=FONT, size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def set_paragraph_left_border(paragraph, color="8C8C8C", size=18):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def style_builtin_heading(style_name, size_pt, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            page_break_before=False, space_before=18, space_after=10,
                            color=(0, 0, 0)):
    st = doc.styles[style_name]
    st.font.name = FONT
    st.font.size = Pt(size_pt)
    st.font.bold = True
    st.font.italic = False
    st.font.color.rgb = RGBColor(*color)
    pf = st.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.page_break_before = page_break_before
    pf.keep_with_next = True
    rpr = st.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)


# stiluri built-in Word, refolosite pentru ca Word sa poata genera automat
# cuprinsul (Table of Contents) din ele
style_builtin_heading("Heading 1", 18, alignment=WD_ALIGN_PARAGRAPH.CENTER, page_break_before=True)
style_builtin_heading("Heading 2", 14, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=16)
style_builtin_heading("Heading 3", 12.5, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(8)
rpr = normal.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.insert(0, rFonts)
rFonts.set(qn("w:ascii"), FONT)
rFonts.set(qn("w:hAnsi"), FONT)
rFonts.set(qn("w:eastAsia"), FONT)

# stil pentru bloc de cod
code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = CODE_FONT
code_style.font.size = Pt(9.5)
code_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
code_style.paragraph_format.space_after = Pt(0)
code_style.paragraph_format.line_spacing = 1.0
code_style.paragraph_format.left_indent = Cm(0.3)
rpr = code_style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.insert(0, rFonts)
rFonts.set(qn("w:ascii"), CODE_FONT)
rFonts.set(qn("w:hAnsi"), CODE_FONT)

# stiluri liste, aliniate la Times New Roman 12
for lst_style in ("List Bullet", "List Number", "List Paragraph"):
    st = doc.styles[lst_style]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.5
    rpr = st.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)

fig_counter = {"n": 0, "chapter": 0}
tab_counter = {"n": 0, "chapter": 0}


# --------------------------------------------------------------- helpers --
def h1(text):
    p = doc.add_heading(text, level=1)
    return p


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def p(text, bold=False, italic=False, align=None, space_after=8):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return para


def p_rich(parts, align=None, space_after=8):
    """parts: lista de (text, bold, italic)"""
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in parts:
        run = para.add_run(text)
        set_run_font(run, bold=bold, italic=italic)
    return para


def bullets(items, numbered=False):
    # numerotare manuala (text literal "1. ", "2. " ...) in loc de stilul
    # nativ "List Number" al Word, ale carui numId-uri se acumuleaza global
    # peste toate listele din document (ar continua 8, 9, 10... intre liste
    # separate, in loc sa reporneasca de la 1 la fiecare lista noua)
    for i, item in enumerate(items, start=1):
        if numbered:
            para = doc.add_paragraph(style="List Paragraph")
            para.paragraph_format.left_indent = Cm(0.63)
            prefix = f"{i}. "
        else:
            para = doc.add_paragraph(style="List Bullet")
            prefix = None
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.5
        if prefix:
            run = para.add_run(prefix + item)
        else:
            run = para.add_run(item)
        set_run_font(run)


def formula(text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    set_run_font(run, italic=True, size=12.5)


def code(text_block, label=None):
    lines = text_block.strip("\n").split("\n")
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(2)
        r = lp.add_run(label)
        set_run_font(r, size=9.5, italic=True, color=(80, 80, 80))
    for i, line in enumerate(lines):
        para = doc.add_paragraph(style="CodeBlock")
        set_paragraph_shading(para, "F2F2F0")
        set_paragraph_left_border(para)
        text = line if line.strip() != "" else " "
        run = para.add_run(text.replace("\t", "    "))
        set_run_font(run, name=CODE_FONT, size=9.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def figure(path, caption_text, width_in=5.6):
    fig_counter["n"] += 1
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    run.add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(f"Figura {fig_counter['chapter']}.{fig_counter['n']} – {caption_text}")
    set_run_font(r, size=10.5, italic=True)


def new_chapter(n):
    fig_counter["chapter"] = n
    fig_counter["n"] = 0
    tab_counter["chapter"] = n
    tab_counter["n"] = 0


def table_simple(headers, rows, caption_text, col_widths_cm=None):
    tab_counter["n"] += 1
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run(f"Tabelul {tab_counter['chapter']}.{tab_counter['n']} – {caption_text}")
    set_run_font(r, size=10.5, italic=True)

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"

    for j, htext in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = ""
        run = cell.paragraphs[0].add_run(htext)
        set_run_font(run, bold=True, size=10.5)
        set_paragraph_shading(cell.paragraphs[0], "E4E2DD")
        if col_widths_cm:
            cell.width = Cm(col_widths_cm[j])

    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10.5)
            if col_widths_cm:
                cell.width = Cm(col_widths_cm[j])

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def add_toc():
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Faceti click dreapta aici si alegeti „Update Field” pentru a genera cuprinsul."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(txt)
    r.append(fld_end)


def add_page_numbers():
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    set_run_font(run, size=10.5)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


add_page_numbers()
section.different_first_page_header_footer = True
# fara numar de pagina pe coperta

# ============================================================ PAGINA DE TITLU
for _ in range(3):
    doc.add_paragraph()

p("[Universitatea …]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p("[Facultatea …]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p("[Programul de studii de master …]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

p_rich(
    [("LUCRARE DE DISERTAȚIE", True, False)],
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30,
)
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_after = 20
r = tp.add_run(
    "Sistem IoT pentru monitorizarea și predicția stării unui filtru de apă. "
    "Arhitectură edge-to-cloud cu evaluarea impactului rețelelor 5G asupra "
    "livrării datelor"
)
set_run_font(r, size=16, bold=True)

for _ in range(4):
    doc.add_paragraph()

info = doc.add_table(rows=2, cols=2)
info.autofit = True
cells = [
    ("Coordonator științific:", "[Nume prenume, titlu didactic]"),
    ("Absolvent:", "Daniel Ailenei"),
]
for i, (label, value) in enumerate(cells):
    c0 = info.cell(i, 0)
    c0.text = ""
    r0 = c0.paragraphs[0].add_run(label)
    set_run_font(r0, bold=True)
    c1 = info.cell(i, 1)
    c1.text = ""
    r1 = c1.paragraphs[0].add_run(value)
    set_run_font(r1)

for _ in range(2):
    doc.add_paragraph()

p("2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

para = doc.add_paragraph()
para.paragraph_format.page_break_before = True

# ============================================================ CUPRINS
h1_toc = doc.add_paragraph()
h1_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h1_toc.add_run("CUPRINS")
set_run_font(r, size=18, bold=True)
h1_toc.paragraph_format.space_after = 20

add_toc()

# ================================================================ CAPITOLUL 1
new_chapter(1)
h1("CAPITOLUL 1. INTRODUCERE")

h2("1.1. Context general")
p(
    "În ultimul deceniu, paradigma Internet of Things (IoT) a schimbat "
    "fundamental modul în care sunt monitorizate și întreținute sistemele "
    "tehnice, prin înlocuirea inspecțiilor manuale, periodice, cu fluxuri "
    "continue de date colectate direct de la echipament. Această tranziție "
    "este vizibilă mai ales în domeniul mentenanței predictive (predictive "
    "maintenance), unde decizia de intervenție nu mai este luată pe baza "
    "unui calendar fix, ci pe baza stării reale, măsurate, a echipamentului."
)
p(
    "Filtrele de apă reprezintă un exemplu tipic de componentă supusă unei "
    "degradări progresive și greu observabile direct de către utilizator: "
    "pe măsură ce impuritățile se acumulează în interiorul materialului "
    "filtrant, porii acestuia se colmatează treptat, ceea ce determină o "
    "creștere a presiunii diferențiale dintre intrarea și ieșirea "
    "filtrului, o scădere a debitului de apă disponibil și, în cazurile "
    "avansate, o scădere a calității apei filtrate (creșterea turbidității). "
    "În absența unei monitorizări active, înlocuirea filtrului se face de "
    "regulă după un interval de timp fix, stabilit empiric de producător, "
    "fără a ține cont de condițiile reale de utilizare (calitatea apei de "
    "la sursă, volumul consumat, presiunea de rețea etc.). Această abordare "
    "conduce fie la înlocuiri premature, cu costuri suplimentare inutile, "
    "fie la înlocuiri întârziate, cu impact direct asupra calității apei "
    "consumate."
)
p(
    "În paralel, dezvoltarea rețelelor mobile de generația a cincea (5G) "
    "aduce, pe lângă viteze de transfer superioare, o reducere semnificativă "
    "a latenței de comunicație și o capacitate mult mai mare de conectare a "
    "unui număr mare de dispozitive IoT simultan. Aceste caracteristici sunt "
    "relevante direct pentru sisteme de monitorizare de tip senzor-cloud, "
    "unde timpul de livrare a datelor și stabilitatea acestuia (jitter-ul) "
    "influențează atât promptitudinea reacției sistemului, cât și calitatea "
    "predicțiilor realizate pe baza datelor colectate."
)

h2("1.2. Fundamentare teoretică")
p(
    "Înainte de a detalia arhitectura sistemului propriu-zis (Capitolul 2), "
    "această secțiune introduce pe scurt două concepte teoretice pe care "
    "se sprijină întreaga soluție: modelul de comunicație client-server, "
    "care stă la baza schimbului de date dintre componentele sistemului, "
    "și modelul relațional al bazelor de date, folosit ca punct de "
    "referință pentru a justifica, prin contrast, alegerea unei baze de "
    "date de tip serie-de-timp (secțiunea 2.2.3)."
)

h3("1.2.1. Internetul și modelul de comunicație client-server")
p(
    "Internetul a apărut, în anii 1960-1970, din nevoia militară și "
    "academică americană de a construi o rețea de comunicație "
    "descentralizată, capabilă să funcționeze chiar dacă o parte din "
    "infrastructură ar fi fost distrusă — primele conexiuni experimentale "
    "între calculatoare aflate în locații diferite datează din 1965. De "
    "atunci, Internetul a evoluat într-o rețea globală de rețele, iar "
    "World Wide Web (WWW) — sistemul de documente și aplicații interconectate "
    "prin hyperlink-uri, accesibil prin protocolul HTTP — a devenit modul "
    "predominant prin care aplicațiile comunică peste Internet."
)
p(
    "Majoritatea aplicațiilor web moderne, inclusiv componentele acestui "
    "sistem, se bazează pe modelul client-server: o aplicație client "
    "trimite o cerere unei aplicații server, care rămâne în așteptare "
    "până primește cereri, le procesează și returnează un răspuns. "
    "Fiecare resursă este identificată printr-un URL (Uniform Resource "
    "Locator), care specifică protocolul de comunicație, adresa gazdei și "
    "calea către resursă — de exemplu, http://localhost:8000/predict, "
    "folosit pentru a interoga modelul de predicție al backend-ului "
    "(secțiunea 3.3)."
)
p(
    "În sistemul descris în această lucrare, modelul client-server se "
    "regăsește direct în comunicația dintre browser (sau clientul curl) "
    "și API-ul REST al backend-ului, respectiv între Grafana și InfluxDB. "
    "Comunicația dintre senzorul virtual și backend, prin MQTT, urmează "
    "însă un model diferit — publish/subscribe (detaliat în secțiunea "
    "2.5.1) — în care sursa datelor (senzorul) nu comunică direct cu "
    "destinatarul (backend-ul), ci printr-un intermediar (broker-ul "
    "Mosquitto). Distincția este importantă: sistemul combină, deliberat, "
    "ambele modele de comunicație, fiecare potrivit unui tip diferit de "
    "interacțiune — cereri punctuale, sincrone (client-server) față de "
    "fluxuri continue de evenimente, asincrone (publish/subscribe)."
)

h3("1.2.2. Bazele de date relaționale")
p(
    "Bazele de date relaționale organizează informația sub formă de "
    "tabele compuse din rânduri și coloane, fiecare rând fiind identificat "
    "printr-o cheie (de regulă, un identificator unic), iar relațiile "
    "dintre tabele fiind exprimate prin chei străine — de exemplu, un "
    "tabel de comenzi care referă, prin ID-ul clientului, un tabel "
    "separat de clienți. Acest model, standardizat prin limbajul SQL, a "
    "devenit dominant în industrie datorită capacității sale de a elimina "
    "duplicarea datelor și de a impune reguli stricte de consistență "
    "(constrângeri de integritate)."
)
p(
    "Modelul relațional este optimizat pentru date structurate, cu "
    "relații complexe între entități, și interogări care combină "
    "informații din mai multe tabele. Datele colectate de sistemul "
    "descris în această lucrare au însă o structură mult mai simplă — "
    "o singură entitate (citirea de senzor), fără relații între "
    "înregistrări, dar cu un volum foarte mare de scrieri secvențiale, "
    "ordonate strict temporal. Această diferență de tipar de acces este "
    "motivul pentru care, în locul unei baze de date relaționale, "
    "sistemul folosește o bază de date specializată pentru serii de timp "
    "(InfluxDB) — alegere justificată în detaliu în secțiunile 2.2.3 și "
    "2.5.2."
)

h2("1.3. Motivația alegerii temei")
p(
    "Alegerea acestei teme a fost motivată de dorința de a construi, într-un "
    "cadru practic și reproductibil, un sistem IoT complet — de la nivelul "
    "senzorului, până la stocarea datelor, predicția stării echipamentului "
    "și vizualizarea rezultatelor — care să integreze, într-un singur studiu "
    "de caz coerent, mai multe direcții tehnologice de actualitate: "
    "comunicația de tip publish/subscribe specifică IoT (MQTT), stocarea "
    "seriilor de timp la scară (InfluxDB), predicția bazată pe modele "
    "statistice (regresie liniară aplicată pe un fenomen fizic cu evoluție "
    "exponențială), vizualizarea datelor în timp real (Grafana) și "
    "evaluarea impactului condițiilor de rețea asupra livrării datelor, "
    "prin simulare de rețea (ns-3)."
)
p(
    "Un aspect important al motivației a fost și constrângerea practică de "
    "a nu dispune, în cadrul acestei lucrări, de un filtru de apă real, "
    "instrumentat cu senzori fizici, și nici de acces la o infrastructură "
    "5G funcțională pentru testare directă. În loc să limiteze scopul "
    "lucrării, această constrângere a condus la o decizie metodologică "
    "deliberată și asumată: înlocuirea componentelor fizice cu simulări "
    "riguroase — un senzor virtual, guvernat de un model matematic de "
    "degradare cu bază fizică plauzibilă, respectiv o simulare de rețea "
    "realizată cu ns-3, un simulator de rețea consacrat în mediul academic. "
    "Această abordare permite obținerea unor rezultate reproductibile, "
    "controlabile și comparabile între scenarii diferite — un avantaj greu "
    "de atins cu echipamente fizice reale, unde condițiile experimentale nu "
    "pot fi variate și repetate în mod identic."
)

h2("1.4. Obiectivele lucrării")
p("Obiectivele urmărite prin realizarea prezentei lucrări sunt următoarele:")
bullets(
    [
        "Proiectarea unei arhitecturi IoT complete, de tip edge-to-cloud, "
        "pentru monitorizarea stării unui filtru de apă, cu separarea clară "
        "a responsabilităților pe niveluri (achiziție, transport, procesare, "
        "stocare, vizualizare).",
        "Implementarea unui senzor virtual care simulează, pe baza unui "
        "model matematic de degradare cu evoluție exponențială, "
        "comportamentul unui filtru real (presiune diferențială, debit, "
        "turbiditate).",
        "Colectarea, transportul (prin protocolul MQTT) și stocarea "
        "persistentă a datelor de senzor într-o bază de date de tip "
        "serie-de-timp (InfluxDB).",
        "Dezvoltarea unui model predictiv, bazat pe regresia liniară "
        "aplicată logaritmului presiunii diferențiale, pentru estimarea "
        "numărului de zile rămase până la înfundarea completă a filtrului.",
        "Vizualizarea în timp real a stării filtrului printr-un dashboard "
        "interactiv (Grafana), provizionat automat.",
        "Simularea, cu ajutorul simulatorului de rețea ns-3, a impactului "
        "condițiilor de rețea (o rețea rapidă, neîncărcată, comparativ cu "
        "o rețea congestionată, cu trafic de fond în rafale) asupra "
        "latenței și jitter-ului de livrare a datelor senzorului.",
        "Containerizarea și orchestrarea componentelor sistemului cu "
        "Docker Compose, pentru reproductibilitate, izolare și ușurință de "
        "implementare pe orice sistem de calcul.",
    ],
    numbered=True,
)

h2("1.5. Structura lucrării")
p(
    "Lucrarea este organizată în capitole care urmăresc, în ordine, "
    "proiectarea și implementarea sistemului descris anterior."
)
p_rich(
    [
        ("Capitolul 2", True, False),
        (
            " prezintă arhitectura de ansamblu a sistemului, justifică "
            "alegerile tehnologice realizate pentru fiecare componentă, "
            "descrie fluxul complet al datelor prin sistem și structura "
            "proiectului la nivel de fișiere și directoare.",
            False, False,
        ),
    ]
)
p_rich(
    [
        ("Capitolul 3", True, False),
        (
            " detaliază, componentă cu componentă, implementarea efectivă "
            "a sistemului — senzorul virtual și modelul matematic de "
            "degradare, broker-ul de mesagerie MQTT, backend-ul aplicației "
            "(subscriber MQTT, scriere/interogare a bazei de date, modelul "
            "de predicție), baza de date de tip serie-de-timp, dashboard-ul "
            "de vizualizare, orchestrarea cu Docker Compose, precum și "
            "simularea de rețea realizată cu ns-3 — însoțită de extrase de "
            "cod comentate și explicații ale logicii de funcționare.",
            False, False,
        ),
    ]
)
p_rich(
    [
        ("Capitolul 4", True, False),
        (
            " descrie metodologia de testare a sistemului, rezultatele "
            "experimentale obținute în urma rulărilor efective — inclusiv "
            "verificarea endpoint-urilor API, a dashboard-ului și a "
            "integrării cu simularea de rețea — precum și principalele "
            "dificultăți tehnice întâmpinate în procesul de implementare, "
            "împreună cu soluțiile adoptate.",
            False, False,
        ),
    ]
)
p(
    "Concluziile generale ale lucrării, limitările abordării alese și "
    "direcțiile posibile de continuare vor face obiectul unui capitol "
    "final, ce va continua această structură într-o etapă ulterioară a "
    "lucrării."
)

# ================================================================ CAPITOLUL 2
new_chapter(2)
h1("CAPITOLUL 2. ARHITECTURA SISTEMULUI ȘI PLANUL APLICAȚIEI")

h2("2.1. Prezentare generală a arhitecturii")
p(
    "Sistemul proiectat urmează o arhitectură de tip edge-to-cloud, "
    "structurată pe cinci niveluri funcționale distincte, fiecare "
    "implementat printr-o componentă software dedicată: nivelul de "
    "achiziție (senzorul virtual), nivelul de transport (broker-ul de "
    "mesagerie MQTT), nivelul de procesare și aplicație (backend-ul "
    "FastAPI, care include și modelul de predicție), nivelul de "
    "persistență (baza de date de tip serie-de-timp InfluxDB) și nivelul "
    "de prezentare (dashboard-ul Grafana). Un al șaselea nivel, opțional, "
    "este reprezentat de simularea de rețea realizată cu ns-3, folosită "
    "pentru a evalua impactul condițiilor de rețea asupra livrării "
    "datelor, fără a fi integrată live în fluxul de execuție al "
    "sistemului."
)
p(
    "Separarea pe niveluri are un dublu rol: pe de o parte, reflectă "
    "arhitectura tipică a unui sistem IoT real (unde senzorul, rețeaua de "
    "transport, platforma cloud și interfața de vizualizare sunt, în mod "
    "obișnuit, componente și chiar furnizori diferiți), iar pe de altă "
    "parte permite testarea și înlocuirea independentă a fiecărei "
    "componente — de exemplu, senzorul virtual ar putea fi înlocuit cu "
    "unul fizic, fără nicio modificare a restului sistemului, atât timp "
    "cât acesta publică date în același format, pe același topic MQTT."
)
p(
    "Patru dintre cele cinci componente principale (broker-ul MQTT, "
    "backend-ul, baza de date și dashboard-ul) rulează izolat, în "
    "containere Docker, orchestrate printr-un singur fișier "
    "docker-compose.yml, pornite cu o singură comandă. Senzorul virtual "
    "rulează nativ, direct cu interpretorul Python instalat pe sistemul "
    "gazdă — o decizie explicată în detaliu în secțiunea 2.2 — pentru a "
    "permite modificarea și repornirea sa rapidă în timpul dezvoltării, "
    "fără reconstruirea vreunei imagini Docker."
)

figure(ARCH_DIAGRAM, "Arhitectura de ansamblu a sistemului Water Filter Monitor", width_in=6.2)

p(
    "Fluxul de bază, ilustrat în Figura 2.1, este următorul: senzorul "
    "virtual publică periodic o citire (presiune, debit, turbiditate) pe "
    "un topic MQTT; broker-ul Mosquitto transportă mesajul către "
    "backend-ul FastAPI, abonat la același topic; backend-ul scrie fiecare "
    "citire în InfluxDB și, în paralel, expune un API REST propriu "
    "(/latest, /history, /predict); Grafana interoghează direct InfluxDB, "
    "folosind limbajul Flux, pentru a desena grafice actualizate la "
    "fiecare câteva secunde. Componenta ns-3, marcată punctat în diagramă, "
    "poate alimenta senzorul virtual cu latențe realiste, extrase dintr-o "
    "simulare de rețea rulată separat."
)

h2("2.2. Stack-ul tehnologic și justificarea alegerilor")
p(
    "Tabelul 2.1 sintetizează tehnologiile alese pentru fiecare componentă "
    "a sistemului, împreună cu rolul lor în arhitectura de ansamblu."
)

table_simple(
    ["Componentă", "Tehnologie", "Rol"],
    [
        ["Senzor virtual", "Python 3.12, paho-mqtt", "Simulează degradarea filtrului, publică pe MQTT"],
        ["Broker mesagerie", "Eclipse Mosquitto 2", "Transport MQTT, senzor → backend"],
        ["Backend", "FastAPI, influxdb-client, scikit-learn", "API REST, scriere date, predicție ML"],
        ["Bază de date", "InfluxDB 2.7", "Stocare serie-de-timp a citirilor"],
        ["Vizualizare", "Grafana", "Dashboard live, provizionat automat"],
        ["Orchestrare", "Docker Compose", "Pornire/oprire infrastructură cu o comandă"],
        ["Simulare rețea (opțional)", "ns-3, WSL2", "Latențe realiste, aplicate ca delay"],
    ],
    "Stack-ul tehnologic al sistemului",
    col_widths_cm=[4.0, 5.5, 6.5],
)

h3("2.2.1. Senzor virtual, nu hardware real")
p(
    "Utilizarea unui senzor virtual, în locul unui filtru fizic "
    "instrumentat, a fost o decizie deliberată, motivată de trei "
    "considerente: controlul complet asupra scenariilor de degradare "
    "(posibilitatea de a varia rata de colmatare, condițiile de rețea "
    "etc.), reproductibilitatea perfectă între rulări succesive și "
    "posibilitatea de a accelera timpul de simulare — o colmatare "
    "completă, care într-un sistem real ar dura luni de zile, poate fi "
    "observată în câteva ore, printr-un factor de accelerare configurabil "
    "(time_acceleration)."
)

h3("2.2.2. MQTT, nu HTTP direct de la senzor")
p(
    "MQTT (Message Queuing Telemetry Transport) este protocolul standard "
    "de facto pentru comunicația IoT, bazat pe modelul publish/subscribe: "
    "senzorii publică mesaje pe un topic, fără a cunoaște cine le "
    "consumă, iar consumatorii se abonează la topicele de interes. Spre "
    "deosebire de un apel HTTP direct (care ar necesita ca senzorul să "
    "cunoască adresa exactă a backend-ului și să mențină o conexiune "
    "sincronă), MQTT oferă un overhead minim de protocol, reconectare "
    "automată și decuplare completă între producător și consumator — "
    "caracteristici relevante mai ales în contextul comparației cu "
    "latențele de rețea simulate în Capitolul 3, secțiunea 3.7."
)

h3("2.2.3. InfluxDB, nu o bază de date relațională")
p(
    "Datele colectate sunt, prin natura lor, serii de timp: citiri "
    "numerice, asociate unui moment exact, cu frecvență relativ ridicată "
    "de scriere și interogări realizate aproape exclusiv pe intervale "
    "temporale. InfluxDB este o bază de date specializată pentru acest "
    "tip de date, oferind compresie eficientă, politici de retenție "
    "configurabile și un limbaj de interogare (Flux) optimizat pentru "
    "agregări temporale — caracteristici pe care o bază de date "
    "relațională generică (de exemplu PostgreSQL sau MySQL) le-ar oferi "
    "doar prin extensii sau prin modelare suplimentară, cu o eficiență "
    "mai scăzută la volume mari de date."
)

h3("2.2.4. Regresie liniară, nu o rețea neuronală, pentru predicție")
p(
    "Modelul fizic de degradare a filtrului, detaliat în secțiunea 3.1, "
    "este explicit exponențial. Această proprietate permite liniarizarea "
    "problemei prin logaritmare, transformând predicția într-o regresie "
    "liniară simplă — o soluție rapidă, ieftină din punct de vedere "
    "computațional și, mai important, interpretabilă: coeficientul de "
    "regresie are o semnificație fizică directă (rata de degradare), iar "
    "calitatea potrivirii poate fi cuantificată printr-un indicator "
    "statistic uzual (R²). O rețea neuronală ar putea, teoretic, învăța "
    "aceeași relație, însă ar necesita un volum de date de antrenare mult "
    "mai mare, ar fi mai greu de interpretat și ar introduce un risc de "
    "supra-ajustare nejustificat pentru un fenomen fizic deja cunoscut "
    "analitic. Extinderea către modele neliniare (regresie polinomială, "
    "rețele recurente) rămâne o direcție de dezvoltare ulterioară, "
    "relevantă în special pentru scenarii de degradare mai complexe, "
    "neexponențiale."
)

h3("2.2.5. Docker Compose pentru infrastructură, Python nativ pentru senzor")
p(
    "Componentele de infrastructură (broker, bază de date, backend, "
    "dashboard) beneficiază de izolare, reproductibilitate și pornire "
    "simultană, motiv pentru care sunt orchestrate prin Docker Compose. "
    "Senzorul virtual, fiind componenta cel mai frecvent modificată pe "
    "durata dezvoltării (ajustarea parametrilor modelului de degradare, "
    "testarea de scenarii de rețea diferite), rulează nativ, direct cu "
    "Python, evitând reconstruirea unei imagini Docker la fiecare "
    "modificare de cod."
)

h2("2.3. Fluxul de date prin sistem")
p(
    "Parcursul complet al unei citiri, de la generarea ei de către "
    "senzorul virtual până la afișarea pe dashboard, urmează pașii "
    "descriși mai jos."
)
bullets(
    [
        "Senzorul virtual calculează o citire nouă (presiune, debit, "
        "turbiditate), pe baza modelului matematic de degradare, la "
        "fiecare interval configurat (implicit 5 secunde).",
        "Opțional, senzorul așteaptă un interval de timp (delay) egal cu "
        "latența curentă, citită ciclic dintr-un fișier CSV generat "
        "printr-o simulare de rețea ns-3, înainte de a publica citirea.",
        "Citirea este publicată, sub formă de mesaj JSON, pe topicul MQTT "
        "home/water/filter, către broker-ul Mosquitto.",
        "Backend-ul FastAPI, abonat la același topic, primește mesajul "
        "prin intermediul unui callback asincron și îl scrie, ca punct de "
        "date, în InfluxDB (measurement filter_reading).",
        "Grafana interoghează direct InfluxDB, folosind limbajul Flux, și "
        "actualizează cele patru panouri ale dashboard-ului la fiecare "
        "cinci secunde.",
        "Independent, orice client extern (browser, script, aplicație "
        "terță) poate interoga backend-ul prin API-ul REST expus "
        "(/latest, /history, /predict), pentru a obține ultima citire, "
        "istoricul sau o predicție a numărului de zile rămase până la "
        "înfundare.",
    ],
    numbered=True,
)

h2("2.4. Structura proiectului și planul de implementare")
p(
    "Proiectul este organizat pe directoare, fiecare corespunzând unei "
    "componente din arhitectura descrisă anterior. Structura completă a "
    "proiectului este prezentată mai jos."
)

code(
    """water-filter-monitor/
├── docker-compose.yml
├── sensor/                     # senzor virtual (Python, MQTT publisher)
│   ├── virtual_sensor.py
│   ├── filter_model.py         # modelul matematic de degradare
│   ├── config.yaml
│   └── requirements.txt
├── backend/                    # FastAPI: subscriber MQTT, InfluxDB, predictie ML
│   ├── main.py
│   ├── mqtt_subscriber.py
│   ├── db_writer.py
│   ├── ml_model.py
│   ├── requirements.txt
│   └── Dockerfile
├── network_sim/                 # simulare ns-3 (rulata separat, WSL2)
│   ├── ns3_config/
│   │   └── iot_water_filter_scenario.cc
│   ├── xml_to_csv.py
│   ├── plot_latency_comparison.py
│   ├── latency_output.csv
│   └── latency_congestionat.csv
├── grafana/provisioning/         # datasource + dashboard, provizionate automat
│   ├── datasources/influxdb.yaml
│   └── dashboards/water_filter_dashboard.json
├── mosquitto/config/
│   └── mosquitto.conf
└── docs/
    └── arhitectura.md""",
    label="Structura de directoare a proiectului",
)

p(
    "Orchestrarea componentelor este realizată printr-un singur fișier "
    "docker-compose.yml, care definește patru servicii (mosquitto, "
    "influxdb, backend, grafana), rețeaua internă comună acestora și "
    "porturile expuse către sistemul gazdă. Tabelul 2.2 sintetizează "
    "porturile și dependințele fiecărui serviciu."
)

table_simple(
    ["Serviciu", "Imagine Docker", "Port expus", "Depinde de"],
    [
        ["mosquitto", "eclipse-mosquitto:2", "1883", "—"],
        ["influxdb", "influxdb:2.7", "8086", "—"],
        ["backend", "construit local (Dockerfile)", "8000", "mosquitto, influxdb"],
        ["grafana", "grafana/grafana:latest", "3000", "influxdb"],
    ],
    "Serviciile definite in docker-compose.yml",
    col_widths_cm=[3.5, 5.5, 3.0, 4.0],
)

h2("2.5. Concepte fundamentale ale sistemului")
p(
    "Pentru o înțelegere completă a arhitecturii descrise anterior, "
    "această secțiune introduce, pe scurt, conceptele teoretice pe care "
    "se sprijină fiecare componentă majoră a sistemului: protocolul "
    "MQTT, bazele de date de tip serie-de-timp, regresia statistică, "
    "containerizarea și rețelele mobile de generația a cincea."
)

h3("2.5.1. Protocolul MQTT")
p(
    "MQTT (Message Queuing Telemetry Transport) este un protocol de "
    "nivel aplicație, construit peste TCP/IP, proiectat special pentru "
    "dispozitive cu resurse limitate și rețele cu lățime de bandă redusă "
    "sau nesigure — caracteristici tipice mediului IoT. Funcționează pe "
    "un model publish/subscribe, mediat de un broker central: un client "
    "(în cazul de față, senzorul virtual) publică mesaje pe un topic "
    "(un șir ierarhic de tip home/water/filter), fără a avea nicio "
    "informație despre eventualii destinatari, iar orice alt client "
    "(backend-ul) se poate abona la același topic pentru a primi acele "
    "mesaje. Protocolul definește trei niveluri de calitate a "
    "serviciului (QoS): nivelul 0 (livrare „cel mult o dată”, fără "
    "confirmare), nivelul 1 (livrare „cel puțin o dată”, cu "
    "posibilitate de duplicare) și nivelul 2 (livrare „exact o dată”, "
    "cu cel mai mare overhead de protocol). Implementarea din această "
    "lucrare folosește nivelul implicit al bibliotecii paho-mqtt (QoS "
    "0), adecvat pentru citiri periodice de senzor, unde pierderea "
    "ocazională a unui singur mesaj nu afectează semnificativ predicția "
    "pe termen lung."
)

h3("2.5.2. Baze de date de tip serie-de-timp")
p(
    "O bază de date de tip serie-de-timp (time-series database) este "
    "optimizată pentru stocarea și interogarea unor date structurate ca "
    "perechi (marcaj temporal, valoare), scrise, de regulă, în ordine "
    "cronologică, aproape strict crescătoare, și interogate predominant "
    "pe intervale de timp (de exemplu, „ultimele 24 de ore”). Spre "
    "deosebire de o bază de date relațională, unde fiecare rând ar "
    "trebui identificat printr-o cheie primară artificială și indexat "
    "separat pentru interogări temporale — o bază de date de tip "
    "serie-de-timp organizează datele intern pe blocuri temporale "
    "(shards), aplică algoritmi de compresie specifici seriilor "
    "numerice și oferă funcții de agregare temporală (medie, sumă, "
    "ultimul element etc.) ca operații native ale limbajului de "
    "interogare, nu ca extensii adăugate ulterior."
)

h3("2.5.3. Regresia liniară și liniarizarea unui fenomen exponențial")
p(
    "Regresia liniară este o metodă statistică ce estimează parametrii "
    "unei relații liniare (y = a·x + b) care se potrivește cel mai bine "
    "unui set de puncte observate, minimizând suma pătratelor "
    "diferențelor dintre valorile prezise și cele reale (metoda celor "
    "mai mici pătrate). Deși relația dintre presiunea diferențială și "
    "timp este exponențială, nu liniară, aplicarea logaritmului natural "
    "asupra presiunii transformă problema într-una liniară — o tehnică "
    "standard, numită liniarizare, folosită frecvent atunci când "
    "fenomenul studiat are o formă funcțională cunoscută analitic. "
    "Calitatea unei regresii se evaluează, de regulă, prin coeficientul "
    "de determinare R², care exprimă proporția din variația datelor "
    "explicată de model (o valoare apropiată de 1 indică o potrivire "
    "foarte bună)."
)

h3("2.5.4. Containerizarea și Docker")
p(
    "Un container este o unitate de software care împachetează codul "
    "unei aplicații împreună cu toate dependințele sale (biblioteci, "
    "interpretor, variabile de mediu), izolată la nivel de proces față "
    "de sistemul de operare gazdă, dar fără suprapunerea de resurse "
    "specifică unei mașini virtuale complete. Docker este platforma "
    "software care standardizează construirea (imagini), distribuirea și "
    "rularea containerelor. Docker Compose extinde acest model pentru "
    "aplicații formate din mai multe servicii interconectate (precum "
    "sistemul descris în această lucrare), permițând definirea "
    "declarativă, într-un singur fișier, a tuturor containerelor, "
    "rețelelor și volumelor necesare, pornite și oprite atomic, ca un "
    "singur ansamblu."
)

h3("2.5.5. Rețelele mobile 5G și relevanța lor pentru IoT")
p(
    "Generația a cincea de rețele mobile (5G) aduce, față de generația "
    "precedentă (4G/LTE), trei clase de îmbunătățiri relevante pentru "
    "aplicații IoT: viteze de transfer semnificativ mai mari (eMBB — "
    "enhanced Mobile Broadband), latențe foarte reduse și predictibile "
    "(URLLC — Ultra-Reliable Low-Latency Communications) și capacitate "
    "mult mai mare de conectare simultană a unui număr foarte mare de "
    "dispozitive pe unitatea de suprafață (mMTC — massive Machine-Type "
    "Communications). Pentru un sistem de monitorizare precum cel "
    "descris în această lucrare, componenta URLLC este cea mai "
    "relevantă: o latență mică și stabilă (jitter redus) înseamnă că "
    "datele de senzor ajung la platforma cloud rapid și previzibil, "
    "ceea ce este cu atât mai important cu cât sistemul ar fi extins să "
    "reacționeze automat (de exemplu, printr-o alertă) la o schimbare "
    "bruscă a stării filtrului. Simularea de rețea descrisă în secțiunea "
    "3.7 evaluează exact acest aspect — diferența de latență și jitter "
    "între o rețea rapidă, neîncărcată (analoagă unei conexiuni URLLC), "
    "și una congestionată."
)

# ================================================================ CAPITOLUL 3
new_chapter(3)
h1("CAPITOLUL 3. CODUL APLICAȚIEI ȘI EXPLICAȚII")

p(
    "Acest capitol detaliază, componentă cu componentă, implementarea "
    "efectivă a sistemului descris în Capitolul 2. Pentru fiecare "
    "componentă sunt prezentate: rolul funcțional în arhitectura de "
    "ansamblu, extrase de cod reprezentative (însoțite de explicații ale "
    "logicii de funcționare) și, unde este cazul, formulele matematice "
    "care stau la baza implementării."
)

# ---------------------------------------------------------------- 3.1
h2("3.1. Senzorul virtual")
p(
    "Senzorul virtual (directorul sensor/) este componenta responsabilă "
    "cu generarea citirilor simulate și publicarea lor pe MQTT. Este "
    "structurat în două fișiere Python: filter_model.py, care conține "
    "exclusiv logica matematică a degradării filtrului, și "
    "virtual_sensor.py, care implementează bucla de execuție și "
    "comunicația MQTT — o separare care izolează modelul fizic de "
    "detaliile de infrastructură (conexiune MQTT, citire configurație)."
)

h3("3.1.1. Modelul matematic de degradare (filter_model.py)")
p(
    "Comportamentul unui filtru de apă în curs de colmatare este "
    "aproximat printr-un model cu trei mărimi de stare: presiunea "
    "diferențială, debitul și turbiditatea apei filtrate. Pe măsură ce "
    "filtrul reține impurități, porii materialului filtrant se "
    "colmatează, ceea ce determină o creștere aproximativ exponențială a "
    "presiunii diferențiale în timp:"
)
formula("presiune(t) = presiune_bază · e^(k · t)")
p(
    "unde k este constanta de colmatare (clogging_rate), iar t este "
    "timpul scurs, exprimat în ore de funcționare simulată. Debitul de "
    "apă scade invers proporțional cu presiunea diferențială:"
)
formula("debit(t) = debit_bază / (1 + presiune(t))")
p(
    "iar turbiditatea crește liniar, într-un ritm mai lent, odată cu "
    "presiunea. Peste fiecare valoare calculată se adaugă un zgomot "
    "aleator de mică amplitudine, pentru ca citirile generate să semene "
    "cu măsurători reale de senzor, nu cu o funcție matematică perfectă. "
    "Filtrul este considerat înfundat atunci când presiunea diferențială "
    "depășește un prag configurabil (implicit 1,5 bar)."
)
code(
    """class FilterModel:
    def __init__(self, clogging_rate=0.0008, base_pressure=0.2,
                 base_flow=15.0, base_turbidity=0.5,
                 clog_threshold_bar=1.5):
        self.clogging_rate = clogging_rate
        self.base_pressure = base_pressure
        self.base_flow = base_flow
        self.base_turbidity = base_turbidity
        self.clog_threshold_bar = clog_threshold_bar

    def pressure_drop(self, elapsed_hours):
        degradation = self.base_pressure * math.exp(
            self.clogging_rate * elapsed_hours)
        noise = random.uniform(-0.02, 0.02)
        return round(degradation + noise, 3)

    def flow_rate(self, pressure_drop):
        flow = self.base_flow / (1 + pressure_drop)
        noise = random.uniform(-0.3, 0.3)
        return round(max(flow + noise, 0), 2)

    def is_clogged(self, pressure_drop):
        return pressure_drop >= self.clog_threshold_bar""",
    label="Extras din sensor/filter_model.py — clasa FilterModel",
)
p(
    "Deoarece o colmatare completă ar dura, în realitate, luni de zile, "
    "simularea accelerează timpul printr-un factor configurabil "
    "(time_acceleration, implicit 500): o oră reală de execuție a "
    "senzorului corespunde la 500 de ore simulate de funcționare a "
    "filtrului. Cu setările implicite, filtrul se înfundă complet în "
    "aproximativ cinci ore reale de la pornirea senzorului — suficient de "
    "rapid pentru observarea completă a fenomenului într-o singură "
    "sesiune de testare, dar suficient de lent pentru a putea urmări "
    "evoluția graficelor pas cu pas."
)

h3("3.1.2. Bucla de simulare și publicarea MQTT (virtual_sensor.py)")
p(
    "Fișierul virtual_sensor.py implementează bucla principală de "
    "execuție a senzorului: citește parametrii din config.yaml, "
    "instanțiază modelul de degradare și, la fiecare interval configurat, "
    "calculează o citire nouă și o publică pe MQTT."
)
code(
    """while True:
    elapsed_hours = (time.time() - start_time) / 3600 * time_acceleration
    pressure_drop = model.pressure_drop(elapsed_hours)
    flow = model.flow_rate(pressure_drop)
    turbidity = model.turbidity(pressure_drop)
    clogged = model.is_clogged(pressure_drop)

    payload = {
        "timestamp": time.time(),
        "pressure_drop_bar": pressure_drop,
        "flow_rate_lmin": flow,
        "turbidity_ntu": turbidity,
        "is_clogged": clogged,
        "days_remaining_model": model.estimate_days_remaining(
            elapsed_hours, time_acceleration),
    }

    if latency_cycle is not None:
        delay = next(latency_cycle)
        time.sleep(delay)   # delay real, extras din simularea ns-3

    client.publish(mqtt_cfg["topic"], json.dumps(payload))
    time.sleep(publish_interval)""",
    label="Extras din sensor/virtual_sensor.py — bucla principala",
)
p(
    "Elementul notabil al buclei este integrarea opțională cu simularea "
    "de rețea (detaliată în secțiunea 3.7): dacă parametrul "
    "network.use_latency_file este activat în config.yaml, senzorul "
    "citește ciclic un fișier CSV cu latențe reale, obținute dintr-o "
    "simulare ns-3, și așteaptă exact acel interval înainte de fiecare "
    "publicare — reproducând, la nivel aplicativ, efectul unei rețele de "
    "transport reale asupra timpului de livrare a datelor."
)

h3("3.1.3. Fișierul de configurare (config.yaml)")
p(
    "Toți parametrii ajustabili ai senzorului virtual — adresa "
    "broker-ului MQTT, parametrii modelului de degradare, factorul de "
    "accelerare a timpului și integrarea cu simularea de rețea — sunt "
    "centralizați într-un singur fișier YAML, citit la pornire. Această "
    "separare între cod și configurare permite testarea rapidă a mai "
    "multor scenarii (de exemplu, o rată de colmatare mai agresivă, sau "
    "comutarea între cele două fișiere de latențe rezultate din "
    "simularea ns-3) fără nicio modificare a codului sursă."
)
code(
    """mqtt:
  broker: localhost
  port: 1883
  topic: home/water/filter

simulation:
  clogging_rate: 0.0008
  base_pressure_bar: 0.2
  base_flow_lmin: 15.0
  base_turbidity_ntu: 0.5
  clog_threshold_bar: 1.5
  time_acceleration: 500
  publish_interval_seconds: 5

network:
  use_latency_file: true
  latency_file: ../network_sim/latency_congestionat.csv""",
    label="sensor/config.yaml",
)
p(
    "Parametrul time_acceleration merită o mențiune separată: cu "
    "valoarea implicită 500, fiecare oră reală de execuție a senzorului "
    "corespunde la 500 de ore simulate de funcționare a filtrului, ceea "
    "ce înseamnă că un filtru care, în realitate, s-ar înfunda în peste "
    "un an de utilizare continuă, se înfundă complet, în simulare, în "
    "aproximativ cinci ore — un compromis deliberat între realismul "
    "temporal și posibilitatea practică de a observa întregul ciclu de "
    "degradare într-o singură sesiune de testare sau de demonstrație."
)

# ---------------------------------------------------------------- 3.2
h2("3.2. Broker-ul de mesagerie — Mosquitto")
p(
    "Comunicația dintre senzorul virtual și backend este intermediată de "
    "un broker MQTT (Eclipse Mosquitto), rulat într-un container Docker "
    "separat, ale cărui principii de funcționare (model publish/"
    "subscribe, topicuri, nivelurile de calitate a serviciului) au fost "
    "introduse în secțiunea 2.5.1. Această secțiune detaliază "
    "configurația concretă folosită și rolul broker-ului în arhitectura "
    "de ansamblu."
)
p(
    "Senzorul virtual publică fiecare citire pe topicul unic "
    "home/water/filter, fără să cunoască cine o consumă, iar backend-ul "
    "se abonează la același topic, fără să cunoască sursa mesajelor — "
    "această decuplare completă permite, spre exemplu, adăugarea "
    "ulterioară a altor consumatori (un sistem de alertare, un al doilea "
    "backend de arhivare pe termen lung) fără nicio modificare a "
    "senzorului virtual, sau chiar înlocuirea senzorului virtual cu unul "
    "fizic, fără nicio modificare a backend-ului, atât timp cât formatul "
    "mesajului JSON publicat rămâne neschimbat."
)
code(
    """listener 1883
allow_anonymous true

persistence true
persistence_location /mosquitto/data/
log_dest stdout""",
    label="mosquitto/config/mosquitto.conf",
)
p(
    "Directiva listener 1883 configurează broker-ul să asculte conexiuni "
    "TCP simple pe portul standard MQTT (1883, neprotejat prin TLS), iar "
    "allow_anonymous true permite conectarea clienților fără "
    "autentificare — o simplificare adecvată pentru mediul de "
    "dezvoltare/demonstrație al acestei lucrări, dar nepotrivită pentru "
    "un mediu de producție, unde ar fi necesară autentificarea "
    "clienților (utilizator/parolă sau certificate) și, ideal, criptarea "
    "conexiunilor (MQTT peste TLS, pe portul standard 8883). Directivele "
    "persistence asigură salvarea pe disc a stării interne a broker-ului "
    "(sesiuni, mesaje reținute), pentru a supraviețui unei reporniri a "
    "containerului."
)

# ---------------------------------------------------------------- 3.3
h2("3.3. Backend-ul aplicației — FastAPI")
p(
    "Backend-ul (directorul backend/) este componenta centrală a "
    "sistemului, responsabilă cu trei funcții: recepționarea și "
    "persistarea citirilor primite prin MQTT, expunerea unui API REST "
    "pentru interogarea datelor și calcularea predicției numărului de "
    "zile rămase până la înfundarea filtrului. Este structurat în patru "
    "fișiere: main.py (aplicația FastAPI și definirea endpoint-urilor), "
    "mqtt_subscriber.py (subscriber-ul MQTT), db_writer.py (scriere și "
    "interogare InfluxDB) și ml_model.py (modelul de predicție)."
)

h3("3.3.1. Punctul de intrare al aplicației (main.py)")
p(
    "La pornirea aplicației, FastAPI execută o funcție de tip lifespan, "
    "care pornește subscriber-ul MQTT pe un fir de execuție separat, "
    "înainte ca serverul HTTP să înceapă să accepte cereri, și îl oprește "
    "curat la închiderea aplicației."
)
code(
    """@asynccontextmanager
async def lifespan(app: FastAPI):
    subscriber.start()
    yield
    subscriber.stop()
    db_writer.close()

app = FastAPI(title="Water Filter Monitor API", lifespan=lifespan)

@app.get("/latest")
def get_latest():
    if not latest_reading:
        return {"status": "no_data"}
    return latest_reading

@app.get("/history")
def get_history(hours: int = Query(24, ge=1, le=24 * 30)):
    readings = db_writer.get_recent_readings(hours=hours)
    return {"count": len(readings), "readings": readings}

@app.get("/predict")
def predict(hours: int = Query(24, ge=1, le=24 * 30)):
    readings = db_writer.get_recent_readings(hours=hours)
    return predictor.predict_days_remaining(readings)""",
    label="Extras din backend/main.py",
)

h3("3.3.2. Subscriber-ul MQTT (mqtt_subscriber.py)")
p(
    "Clasa MqttSubscriber gestionează conexiunea la broker-ul Mosquitto "
    "și rulează pe firul propriu de execuție (loop_start), pentru a nu "
    "bloca firul principal al serverului FastAPI. La fiecare mesaj "
    "primit, payload-ul JSON este decodificat și scris în InfluxDB."
)
code(
    """def _on_message(self, client, userdata, msg):
    try:
        data = json.loads(msg.payload.decode("utf-8"))
    except (json.JSONDecodeError, UnicodeDecodeError) as e:
        print(f"[mqtt] Mesaj invalid, ignorat: {e}")
        return

    self.db_writer.write_reading(data)
    with self._lock:
        self.latest_reading_ref.update(data)""",
    label="Extras din backend/mqtt_subscriber.py — callback-ul on_message",
)

h3("3.3.3. Scrierea și interogarea InfluxDB (db_writer.py)")
p(
    "Clasa DBWriter încapsulează toate operațiile asupra InfluxDB: "
    "scrierea unei citiri noi ca punct de date și interogarea "
    "istoricului. Fiecare citire este scrisă cu patru câmpuri "
    "(pressure_drop_bar, flow_rate_lmin, turbidity_ntu, is_clogged), în "
    "measurement-ul filter_reading, cu marcaj temporal la nivel de "
    "nanosecundă."
)
code(
    """def write_reading(self, data: dict):
    point = (
        Point("filter_reading")
        .field("pressure_drop_bar", float(data["pressure_drop_bar"]))
        .field("flow_rate_lmin", float(data["flow_rate_lmin"]))
        .field("turbidity_ntu", float(data["turbidity_ntu"]))
        .field("is_clogged", bool(data.get("is_clogged", False)))
        .time(datetime.now(timezone.utc), WritePrecision.NS)
    )
    self.write_api.write(bucket=INFLUX_BUCKET, org=INFLUX_ORG, record=point)

def get_recent_readings(self, hours: int = 24):
    query = f'''
    from(bucket: "{INFLUX_BUCKET}")
      |> range(start: -{hours}h)
      |> filter(fn: (r) => r._measurement == "filter_reading")
      |> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")
      |> sort(columns: ["_time"])
    '''
    tables = self.query_api.query(query, org=INFLUX_ORG)
    ...""",
    label="Extras din backend/db_writer.py",
)
p(
    "Interogarea folosește limbajul Flux, specific InfluxDB 2.x: "
    "operatorul pivot() reorganizează rezultatul astfel încât fiecare "
    "rând să conțină toate cele patru câmpuri ale unei citiri, nu câte un "
    "rând separat per câmp — format mult mai convenabil pentru "
    "prelucrarea ulterioară, atât în modelul de predicție, cât și în "
    "răspunsurile API."
)

h3("3.3.4. Modelul de predicție (ml_model.py)")
p(
    "Deoarece presiunea diferențială crește exponențial în timp "
    "(secțiunea 3.1.1), logaritmul acesteia crește liniar:"
)
formula("ln(presiune(t)) = ln(presiune_bază) + k · t")
p(
    "Clasa FilterPredictor exploatează direct această proprietate: "
    "aplică o regresie liniară (scikit-learn, LinearRegression) pe "
    "perechile (timp scurs, logaritmul presiunii), extrase din istoricul "
    "recent al citirilor din InfluxDB, apoi extrapolează matematic "
    "momentul în care presiunea va atinge pragul de înfundare."
)
code(
    """log_ys = np.log(ys)
model = LinearRegression()
model.fit(xs, log_ys)

slope = model.coef_[0]        # k, rata de degradare
intercept = model.intercept_  # ln(presiune_baza)

# rezolvam exp(slope*t + intercept) = prag => t
t_threshold = (math.log(self.clog_threshold_bar) - intercept) / slope
seconds_remaining = max(t_threshold - latest_t, 0)
days_remaining = seconds_remaining / 86400

r_squared = model.score(xs, log_ys)""",
    label="Extras din backend/ml_model.py — FilterPredictor.predict_days_remaining",
)
p(
    "Rezultatul returnat de endpoint-ul /predict include, pe lângă "
    "numărul estimat de zile rămase, și coeficientul R² al regresiei — "
    "un indicator direct al calității potrivirii modelului pe datele "
    "observate, util pentru a evalua încrederea în predicție. Avantajul "
    "acestei abordări, comparativ cu un model de tip black-box (de "
    "exemplu o rețea neuronală), este că fiecare termen al calculului are "
    "o interpretare fizică directă, ceea ce facilitează validarea și "
    "explicarea rezultatelor."
)

h3("3.3.5. Containerizarea backend-ului (Dockerfile)")
p(
    "Spre deosebire de celelalte trei servicii (mosquitto, influxdb, "
    "grafana), care folosesc imagini Docker publice, gata construite, "
    "backend-ul este singura componentă construită local, pornind de la "
    "codul sursă propriu, printr-un fișier Dockerfile dedicat."
)
code(
    """FROM python:3.11-slim

WORKDIR /app

COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

COPY . .

EXPOSE 8000

CMD ["uvicorn", "main:app", "--host", "0.0.0.0", "--port", "8000"]""",
    label="backend/Dockerfile",
)
p(
    "Imaginea pornește de la o distribuție Python minimală "
    "(python:3.11-slim), instalează întâi doar fișierul de dependințe "
    "(requirements.txt), iar abia apoi copiază restul codului sursă — "
    "ordinea este intenționată: Docker construiește imaginile pe straturi "
    "(layers) și le pune în cache; separând instalarea dependințelor de "
    "copierea codului, o modificare ulterioară a codului sursă (fără "
    "modificarea dependințelor) nu mai declanșează reinstalarea "
    "pachetelor Python la fiecare reconstruire a imaginii, reducând "
    "semnificativ timpul de build în timpul dezvoltării iterative."
)

h3("3.3.6. Sistemul de alertare (alerting.py)")
p(
    "Pentru a transforma sistemul dintr-un simplu instrument de "
    "monitorizare pasivă într-unul care avertizează activ utilizatorul, "
    "a fost adăugat un modul de alertare, integrat direct în fluxul de "
    "procesare a mesajelor MQTT. La fiecare citire primită, backend-ul "
    "calculează procentul de colmatare curent al filtrului — raportul "
    "dintre presiunea diferențială curentă și pragul de înfundare — și "
    "verifică dacă acesta a trecut de unul din cele trei praguri de "
    "alertă predefinite: 80%, 90% și 100% (înfundare completă)."
)
code(
    """THRESHOLDS = [80, 90, 100]
RESET_BELOW_PCT = 50

class AlertManager:
    def __init__(self):
        self._fired = set()  # praguri deja notificate in ciclul curent

    def check_and_notify(self, pressure_drop_bar, clog_threshold_bar):
        pct = (pressure_drop_bar / clog_threshold_bar) * 100

        if pct < RESET_BELOW_PCT and self._fired:
            self.reset()  # filtru inlocuit - repornim pragurile

        for threshold in THRESHOLDS:
            if pct >= threshold and threshold not in self._fired:
                self._fired.add(threshold)
                # trimite pragul rotund (80/90/100), nu raportul brut -
                # vezi sectiunea 4.7 pentru motivul acestei alegeri
                self._send_alert(threshold, pressure_drop_bar, clog_threshold_bar)""",
    label="Extras din backend/alerting.py — AlertManager",
)
p(
    "Un aspect important al proiectării este mecanismul de deduplicare: "
    "fiecare prag este notificat o singură dată per „ciclu” al filtrului, "
    "prin reținerea pragurilor deja declanșate într-o mulțime (_fired). "
    "Fără acest mecanism, senzorul virtual — care publică o citire la "
    "fiecare cinci secunde — ar genera sute de alerte identice pe durata "
    "intervalului în care presiunea rămâne peste un prag deja depășit. "
    "Când presiunea scade brusc sub 50% (semn că filtrul a fost înlocuit "
    "sau senzorul repornit — comportament observat și documentat în "
    "secțiunea 4.3), mulțimea de praguri notificate este resetată, "
    "pentru ca următorul ciclu de colmatare să poată genera din nou "
    "alertele corespunzătoare."
)
p(
    "Fiecare alertă este trimisă simultan pe două canale independente: "
    "email, prin protocolul SMTP (folosind serverul Gmail și o parolă "
    "de aplicație dedicată, nu parola reală a contului), și notificare "
    "push pe telefon, prin serviciul gratuit ntfy.sh — un serviciu "
    "minimalist, care nu necesită autentificare sau cont, unde mesajele "
    "trimise printr-o simplă cerere HTTP POST către un topic (un șir de "
    "caractere ales de utilizator, care funcționează practic ca o parolă "
    "simplă) sunt livrate instant oricărui dispozitiv abonat la acel "
    "topic."
)
code(
    """def _send_push(self, subject, body):
    requests.post(
        NTFY_URL,
        data=body.encode("utf-8"),
        headers={
            "Title": subject.encode("utf-8"),
            "Priority": "urgent" if "INFUNDAT" in subject else "default",
        },
    )""",
    label="Extras din backend/alerting.py — trimiterea notificarii push",
)
p(
    "Credențialele (parola SMTP, topicul ntfy) nu sunt niciodată scrise "
    "în codul sursă sau în fișierele urcate pe repository-ul public al "
    "proiectului — sunt citite exclusiv din variabile de mediu, "
    "furnizate printr-un fișier local (.env.secrets), listat explicit în "
    ".gitignore. Repository-ul conține doar un fișier-șablon "
    "(.env.secrets.example), cu structura așteptată, dar fără valori "
    "reale — o practică standard de gestionare a secretelor în proiecte "
    "open-source."
)

# ---------------------------------------------------------------- 3.4
h2("3.4. Persistența datelor — InfluxDB")
p(
    "InfluxDB este o bază de date specializată pentru serii de timp "
    "(time-series database), organizată pe trei concepte principale: "
    "bucket-ul (echivalentul unei baze de date), measurement-ul "
    "(echivalentul unui tabel — în cazul de față, filter_reading) și "
    "câmpurile (fields), care conțin valorile numerice sau booleene "
    "efective, asociate unui marcaj temporal. Spre deosebire de o bază "
    "de date relațională, InfluxDB este optimizată intern pentru "
    "scrieri secvențiale, ordonate temporal, și pentru interogări care "
    "agregă date pe ferestre de timp — exact tiparul de acces al acestei "
    "aplicații."
)
p(
    "La prima pornire, containerul InfluxDB se auto-configurează, pe "
    "baza variabilelor de mediu definite în docker-compose.yml, cu o "
    "organizație (disertatie), un bucket (water_filter), un utilizator "
    "administrator și un token de acces folosit de backend pentru "
    "autentificare. Interfața web nativă a InfluxDB (disponibilă la "
    "portul 8086) permite, suplimentar față de Grafana, inspectarea "
    "directă a datelor brute, utilă în special pentru depanare."
)
figure(
    os.path.join(HERE, "screenshot_influxdb.jpg"),
    "Interfața Data Explorer a InfluxDB, afișând evoluția reală a "
    "câmpului flow_rate_lmin pe ultimele 6 ore",
    width_in=6.3,
)
p(
    "Captura de mai sus, realizată în timpul testării efective a "
    "sistemului, ilustrează un ciclu complet de funcționare: debitul "
    "scade treptat, pe măsură ce filtrul se colmatează, urmat de o "
    "revenire bruscă la valoarea inițială — moment care corespunde unei "
    "reporniri a senzorului virtual în timpul sesiunii de testare, "
    "confirmând totodată caracterul continuu și persistent al datelor "
    "stocate în InfluxDB, indiferent de repornirile componentelor "
    "amonte."
)
p(
    "Comutatorul „View Raw Data” al Data Explorer-ului înlocuiește "
    "graficul cu tabelul de rezultate brut, așa cum este întors efectiv "
    "de interogarea Flux, înainte de a fi randat grafic. Figura 3.2 "
    "prezintă acest tabel, mărit, pentru lizibilitate."
)
figure(
    os.path.join(HERE, "screenshot_influxdb_table_zoom.png"),
    "Tabelul de date brute din InfluxDB (View Raw Data), mărit pentru "
    "lizibilitate",
    width_in=6.3,
)
p(
    "Coloanele tabelului corespund exact modelului de date al InfluxDB: "
    "table (identificatorul intern al seriei), _measurement și _field "
    "(care măsurătoare și care câmp), _value (valoarea numerică propriu-"
    "zisă), _start / _stop (limitele ferestrei de timp interogate) și "
    "_time (marcajul temporal al fiecărei citiri individuale). Bara de "
    "paginare de la baza tabelului (1, 2, 3 … 96) confirmă volumul mare "
    "de citiri acumulate pe durata sesiunilor de testare din această "
    "lucrare — fiecare pagină corespunzând unui grup de citiri "
    "consecutive."
)

# ---------------------------------------------------------------- 3.5
h2("3.5. Vizualizarea — Grafana")
p(
    "Dashboard-ul Grafana este provizionat complet automat, fără nicio "
    "configurare manuală din interfața web, prin două fișiere de "
    "configurare citite la pornirea containerului."
)

h3("3.5.1. Conectarea automată la sursa de date")
p(
    "Fișierul grafana/provisioning/datasources/influxdb.yaml conectează "
    "Grafana la InfluxDB, fără nicio interacțiune manuală din interfața "
    "web — adresa serviciului, organizația, bucket-ul implicit și "
    "token-ul de acces sunt citite direct din acest fișier la pornirea "
    "containerului."
)
code(
    """datasources:
  - name: InfluxDB
    uid: influxdb-water-filter
    type: influxdb
    access: proxy
    url: http://influxdb:8086
    jsonData:
      version: Flux
      organization: disertatie
      defaultBucket: water_filter
    secureJsonData:
      token: dev-super-secret-token""",
    label="grafana/provisioning/datasources/influxdb.yaml",
)
p(
    "Câmpul uid, fixat explicit la valoarea influxdb-water-filter, "
    "merită o atenție specială: fiecare panou al dashboard-ului "
    "referă sursa de date prin acest identificator. În absența unei "
    "valori fixe, Grafana generează un identificator aleator la fiecare "
    "(re)pornire, iar panourile create anterior, care rețin identifi"
    "catorul vechi, ajung să nu mai găsească sursa de date — exact "
    "problema documentată în secțiunea 4.7, întâlnită și rezolvată în "
    "procesul de dezvoltare al acestei lucrări."
)

h3("3.5.2. Structura dashboard-ului")
p(
    "Fișierul grafana/provisioning/dashboards/water_filter_dashboard.json "
    "definește complet structura dashboard-ului: titlul, intervalul de "
    "reîmprospătare (5 secunde) și cele patru panouri, sintetizate în "
    "Tabelul 3.1."
)
table_simple(
    ["Panou", "Tip", "Câmp InfluxDB afișat"],
    [
        ["Cădere de presiune (bar)", "serie de timp (grafic linie)", "pressure_drop_bar"],
        ["Debit (L/min)", "serie de timp (grafic linie)", "flow_rate_lmin"],
        ["Turbiditate (NTU)", "serie de timp (grafic linie)", "turbidity_ntu"],
        ["Filtru înfundat? (ultima citire)", "indicator text (stat)", "is_clogged (mapat „OK”/„INFUNDAT”)"],
    ],
    "Panourile dashboard-ului Water Filter Monitor",
    col_widths_cm=[5.5, 5.0, 4.0],
)
code(
    """{
  "title": "Cadere de presiune (bar)",
  "type": "timeseries",
  "datasource": { "type": "influxdb", "uid": "influxdb-water-filter" },
  "targets": [{
    "query": "from(bucket: \\"water_filter\\") \\
|> range(start: v.timeRangeStart, stop: v.timeRangeStop) \\
|> filter(fn: (r) => r._measurement == \\"filter_reading\\") \\
|> filter(fn: (r) => r._field == \\"pressure_drop_bar\\")"
  }]
}""",
    label="Extras din grafana/provisioning/dashboards/water_filter_dashboard.json",
)
p(
    "Fiecare panou execută propria interogare Flux, filtrată pe câmpul "
    "corespunzător, în intervalul de timp selectat curent în interfață "
    "(v.timeRangeStart / v.timeRangeStop — variabile injectate automat "
    "de Grafana, sincronizate cu selectorul de interval din colțul "
    "dreapta-sus al dashboard-ului). Panoul de stare a filtrului "
    "folosește, suplimentar, o funcție de mapare a valorilor (value "
    "mapping), care traduce valoarea booleană is_clogged în textul "
    "afișat și culoarea de fundal a panoului."
)

figure(
    os.path.join(HERE, "screenshot_grafana_dashboard.jpg"),
    "Dashboard-ul „Water Filter Monitor”, cu cele patru panouri populate "
    "cu date reale",
    width_in=6.3,
)
p(
    "Captura de mai sus provine dintr-o rulare reală de testare, pe "
    "parcursul căreia se observă clar tendința crescătoare a presiunii "
    "diferențiale și a turbidității, respectiv tendința descrescătoare a "
    "debitului — urmate de o revenire bruscă la valorile inițiale, "
    "corespunzătoare unei reporniri a senzorului virtual în timpul "
    "sesiunii de testare (vizibilă simultan pe toate cele trei grafice, "
    "confirmând consistența datelor între panouri). Panoul de stare "
    "afișează „OK”, presiunea nefiind, la momentul capturii, peste "
    "pragul de înfundare."
)

# ---------------------------------------------------------------- 3.6
h2("3.6. Orchestrarea sistemului — Docker Compose")
p(
    "Cele patru servicii de infrastructură sunt definite într-un singur "
    "fișier docker-compose.yml, care specifică, pentru fiecare serviciu, "
    "imaginea Docker folosită (sau instrucțiunea de construire locală, "
    "în cazul backend-ului), variabilele de mediu, porturile expuse către "
    "sistemul gazdă și volumele de stocare persistentă."
)
code(
    """services:
  mosquitto:
    image: eclipse-mosquitto:2
    ports: ["1883:1883"]
    volumes:
      - ./mosquitto/config:/mosquitto/config

  influxdb:
    image: influxdb:2.7
    ports: ["8086:8086"]
    environment:
      - DOCKER_INFLUXDB_INIT_ORG=disertatie
      - DOCKER_INFLUXDB_INIT_BUCKET=water_filter
    volumes:
      - influxdb-data:/var/lib/influxdb2

  backend:
    build: ./backend
    depends_on: [mosquitto, influxdb]
    environment:
      - MQTT_BROKER=mosquitto
      - INFLUX_URL=http://influxdb:8086
    ports: ["8000:8000"]

  grafana:
    image: grafana/grafana:latest
    depends_on: [influxdb]
    ports: ["3000:3000"]
    volumes:
      - ./grafana/provisioning:/etc/grafana/provisioning""",
    label="Extras din docker-compose.yml",
)
p(
    "Un detaliu important este modul în care serviciile comunică între "
    "ele: toate cele patru containere rulează în aceeași rețea Docker "
    "internă, creată automat, și se identifică unele pe altele după "
    "numele serviciului (de exemplu, backend-ul se conectează la "
    "mosquitto:1883, nu la localhost:1883 — adresare validă doar între "
    "containere). Senzorul virtual, care rulează direct pe sistemul "
    "gazdă și nu într-un container, se conectează la localhost:1883, "
    "posibil datorită publicării explicite a portului către sistemul "
    "gazdă (secțiunea ports din definiția serviciului mosquitto)."
)

# ---------------------------------------------------------------- 3.7
h2("3.7. Simularea de rețea — ns-3")
p(
    "Pentru a evalua impactul condițiilor de rețea asupra timpului de "
    "livrare a datelor senzorului, a fost realizată o simulare de rețea "
    "independentă, folosind ns-3 — un simulator de rețea discret, "
    "consacrat în cercetarea academică din domeniul rețelelor de "
    "calculatoare. Simularea a fost rulată separat, într-un mediu Linux "
    "(WSL2), rezultatele fiind exportate ca fișiere CSV, pe care "
    "senzorul virtual le poate reproduce ca întârziere (delay) reală, "
    "înainte de fiecare publicare MQTT — fără nicio integrare live între "
    "cele două medii de execuție."
)

h3("3.7.1. Topologia simulată")
p(
    "Scenariul simulat modelează șase noduri senzor IoT, conectate "
    "printr-o topologie de tip stea la un nod „AP” (punct de acces), "
    "care la rândul lui este conectat la un nod „server” printr-o "
    "legătură dedicată, denumită backhaul — segmentul de rețea ale cărui "
    "caracteristici (debit, întârziere, dimensiunea cozii) sunt variate "
    "între scenarii, pentru a modela fie o rețea rapidă, neîncărcată "
    "(analoagă unei conexiuni 5G), fie o rețea congestionată. Opțional, "
    "un al șaptelea nod („interferer”) trimite trafic UDP de volum mare, "
    "în rafale aleatoare, peste legătura backhaul, pentru a satura "
    "intermitent legătura și a genera jitter real de congestie."
)
code(
    """PointToPointHelper backhaulLink;
backhaulLink.SetDeviceAttribute("DataRate", StringValue(backhaulDataRate));
backhaulLink.SetChannelAttribute("Delay", StringValue(backhaulDelay));
backhaulLink.SetQueue("ns3::DropTailQueue", "MaxSize",
                       StringValue(std::to_string(backhaulQueueSize) + "p"));

if (enableBackgroundTraffic)
{
    OnOffHelper bgClient("ns3::UdpSocketFactory",
                          InetSocketAddress(serverAddress, backgroundPort));
    bgClient.SetAttribute("DataRate", StringValue(backgroundDataRate));
    bgClient.SetAttribute("OnTime", StringValue(
        "ns3::ExponentialRandomVariable[Mean=" +
        std::to_string(backgroundOnTimeMean) + "]"));
    bgClient.SetAttribute("OffTime", StringValue(
        "ns3::ExponentialRandomVariable[Mean=" +
        std::to_string(backgroundOffTimeMean) + "]"));
    ...
}

FlowMonitorHelper flowmonHelper;
Ptr<FlowMonitor> monitor = flowmonHelper.InstallAll();
monitor->SetAttribute("DelayBinWidth", DoubleValue(0.001));
Simulator::Run();
monitor->SerializeToXmlFile(outputXml, true, true);""",
    label="Extras din network_sim/ns3_config/iot_water_filter_scenario.cc",
)
p(
    "Modulul FlowMonitor, integrat nativ în ns-3, înregistrează automat "
    "statistici de latență pentru fiecare flux de date din simulare, sub "
    "forma unui histogram (numărul de pachete a căror latență a căzut "
    "într-un anumit interval, de exemplu „între 10 și 11 ms, 4 "
    "pachete”), exportat la finalul simulării ca fișier XML."
)

h3("3.7.2. Conversia rezultatelor în format utilizabil de aplicație")
p(
    "Modulul FlowMonitor nu salvează latența fiecărui pachet individual, "
    "ci un histogram: pentru fiecare flux, câte pachete au avut latența "
    "într-un anumit interval de o milisecundă (de exemplu, „între 10 și "
    "11 ms au fost 4 pachete”). Scriptul xml_to_csv.py parcurge fișierul "
    "XML, identifică fluxurile de interes (filtrând după portul de "
    "destinație, implicit 9, pe baza secțiunii Ipv4FlowClassifier — "
    "astfel încât traficul de fond, trimis pe alt port, să nu fie "
    "inclus) și „despachetează” histogramul într-o listă aproximativă de "
    "latențe per pachet: fiecare pachet dintr-un interval primește "
    "valoarea de mijloc a intervalului respectiv, plus un zgomot aleator "
    "de mică amplitudine, ca valorile să nu fie perfect identice."
)
code(
    """def get_sensor_flow_ids(root, port):
    classifier = root.find("Ipv4FlowClassifier")
    flow_ids = set()
    for flow in classifier.findall("Flow"):
        if flow.get("destinationPort") == str(port):
            flow_ids.add(flow.get("flowId"))
    return flow_ids

def extract_delay_bins(xml_path, port=9):
    tree = ET.parse(xml_path)
    root = tree.getroot()
    sensor_flow_ids = get_sensor_flow_ids(root, port)

    bins = []
    for flow in root.find("FlowStats").findall("Flow"):
        if flow.get("flowId") not in sensor_flow_ids:
            continue          # trafic de fond, ignorat
        for bin_el in flow.find("delayHistogram").findall("bin"):
            start = float(bin_el.get("start"))
            width = float(bin_el.get("width"))
            count = int(float(bin_el.get("count")))
            mid_ms = (start + width / 2.0) * 1000.0
            bins.append((mid_ms, count))
    return bins""",
    label="Extras din network_sim/xml_to_csv.py",
)
p(
    "Rezultatul este scris într-un fișier CSV cu formatul "
    "packet_id,latency_ms, citit direct de sensor/virtual_sensor.py. Un "
    "al doilea script, plot_latency_comparison.py, citește ambele "
    "fișiere CSV rezultate (rețea rapidă și rețea congestionată) și "
    "generează graficul comparativ prezentat în Figura 3.4, folosind "
    "biblioteca matplotlib."
)

h3("3.7.3. Rezultate obținute")
p(
    "Au fost rulate două scenarii comparative, fiecare simulând 60 de "
    "secunde de trafic generat de șase senzori. Tabelul 3.2 sintetizează "
    "statisticile obținute."
)
table_simple(
    ["Scenariu", "Minim", "Maxim", "Medie", "Deviație standard"],
    [
        ["Rețea rapidă (50 Mbps, 10 ms, fără trafic de fond)", "11,2 ms", "11,8 ms", "11,5 ms", "≈ 0,15 ms"],
        ["Rețea congestionată (5 Mbps, 40 ms, trafic de fond în rafale)", "41,2 ms", "67,8 ms", "60,5 ms", "≈ 11 ms"],
    ],
    "Statistici de latenta pentru cele doua scenarii simulate",
    col_widths_cm=[6.5, 2.2, 2.2, 2.2, 2.9],
)
figure(
    LATENCY_CHART,
    "Distribuția latenței de livrare a datelor — rețea rapidă vs. rețea congestionată",
    width_in=6.0,
)
p(
    "Diferența dintre cele două scenarii nu este vizibilă doar la nivelul "
    "mediei (de aproximativ cinci ori mai mare în scenariul congestionat), "
    "ci mai ales la nivelul variabilității: deviația standard crește de "
    "la o valoare practic neglijabilă (0,15 ms), în scenariul "
    "neîncărcat, la aproximativ 11 ms, în scenariul congestionat. "
    "Distribuția din acest din urmă caz este vizibil bimodală — un grup "
    "de valori în jurul a 41 ms, corespunzător pachetelor ajunse când "
    "legătura era liberă, și un grup în jurul a 65–68 ms, corespunzător "
    "pachetelor prinse în mijlocul unei rafale de trafic de fond — un "
    "rezultat coerent cu fenomenul fizic de jitter cauzat de o legătură "
    "partajată, suprasolicitată intermitent, și distinct de o simplă "
    "creștere a latenței medii de propagare."
)

# ---------------------------------------------------------------- 3.8
h2("3.8. Sinteza integrării componentelor")
p(
    "Componentele prezentate în acest capitol formează, împreună, "
    "fluxul complet descris în Capitolul 2: senzorul virtual (secțiunea "
    "3.1) generează date pe baza modelului matematic de degradare, "
    "aplicând opțional latențele reale rezultate din simularea de rețea "
    "(secțiunea 3.7); broker-ul MQTT (secțiunea 3.2) transportă aceste "
    "date către backend (secțiunea 3.3), care le persistă în InfluxDB "
    "(secțiunea 3.4) și calculează predicția stării filtrului; Grafana "
    "(secțiunea 3.5) le afișează în timp real, iar întregul ansamblu de "
    "servicii este pornit și gestionat printr-o singură comandă, datorită "
    "orchestrării cu Docker Compose (secțiunea 3.6). Capitolele "
    "următoare ale lucrării vor detalia metodologia de testare a "
    "sistemului, rezultatele experimentale obținute în urma rulărilor "
    "efective și concluziile desprinse din întregul demers."
)

# ================================================================ CAPITOLUL 4
new_chapter(4)
h1("CAPITOLUL 4. TESTAREA SISTEMULUI ȘI REZULTATE EXPERIMENTALE")

p(
    "Acest capitol descrie metodologia folosită pentru verificarea "
    "funcțională a sistemului implementat, rezultatele obținute în urma "
    "rulărilor efective ale acestuia și principalele dificultăți tehnice "
    "întâmpinate în procesul de implementare, împreună cu soluțiile "
    "adoptate. Spre deosebire de o testare unitară, izolată, a codului "
    "sursă, verificarea descrisă aici este de tip end-to-end: sistemul a "
    "fost pornit în întregime, componentă cu componentă, iar "
    "funcționarea corectă a fost confirmată prin observarea directă a "
    "datelor care circulă efectiv prin toate nivelurile arhitecturii, de "
    "la senzor până la dashboard."
)

h2("4.1. Metodologia de testare")
p(
    "Testarea sistemului a urmat o strategie incrementală, de jos în "
    "sus (bottom-up), verificând fiecare nivel al arhitecturii înainte "
    "de a trece la următorul:"
)
bullets(
    [
        "Verificarea pornirii infrastructurii Docker — confirmarea "
        "faptului că toate cele patru servicii (mosquitto, influxdb, "
        "backend, grafana) ajung în starea „Up”, fără erori în log-uri.",
        "Verificarea endpoint-urilor REST ale backend-ului, folosind "
        "clientul de linie de comandă curl, pentru fiecare rută expusă "
        "(/health, /latest, /history, /predict).",
        "Verificarea fluxului complet de date, prin pornirea senzorului "
        "virtual și observarea directă a propagării unei citiri prin "
        "MQTT, InfluxDB, până la afișarea ei pe dashboard-ul Grafana.",
        "Verificarea integrării cu simularea de rețea ns-3, prin "
        "compararea intervalului real dintre publicările succesive ale "
        "senzorului cu delay-ul așteptat, extras din scenariul de "
        "rețea simulat.",
    ],
    numbered=True,
)

h2("4.2. Verificarea pornirii infrastructurii")
p(
    "La pornirea comenzii docker compose up --build, toate cele patru "
    "servicii au fost aduse în stare funcțională, confirmată atât prin "
    "log-urile individuale ale fiecărui container, cât și prin comanda "
    "docker compose ps. Tabelul 4.1 sintetizează rezultatul verificării."
)
table_simple(
    ["Serviciu", "Comportament observat la pornire"],
    [
        ["mosquitto", "Ascultă corect pe portul 1883 (IPv4 și IPv6)"],
        ["influxdb", "Bucket water_filter și organizația disertatie create automat"],
        ["backend", "„Uvicorn running on http://0.0.0.0:8000”, „Application startup complete”"],
        ["grafana", "Datasource InfluxDB și dashboard-ul provizionate automat, server pornit pe portul 3000"],
    ],
    "Verificarea pornirii serviciilor Docker",
    col_widths_cm=[3.0, 11.5],
)

h2("4.3. Verificarea endpoint-urilor API")
p(
    "Fiecare endpoint expus de backend a fost verificat individual, "
    "confirmându-se atât codul de răspuns, cât și structura JSON "
    "returnată. Endpoint-ul /health confirmă disponibilitatea "
    "serviciului:"
)
code(
    """$ curl.exe http://localhost:8000/health
{"status":"ok"}""",
    label="Verificarea endpoint-ului /health",
)
p(
    "Endpoint-ul /latest a fost verificat imediat după pornirea "
    "senzorului virtual, confirmând recepționarea corectă, prin MQTT, a "
    "ultimei citiri publicate:"
)
code(
    """$ curl.exe http://localhost:8000/latest
{"timestamp":1786344815.567231,"pressure_drop_bar":0.202,
 "flow_rate_lmin":12.78,"turbidity_ntu":0.89,
 "is_clogged":false,"days_remaining_model":0.21}""",
    label="Raspuns real al endpoint-ului /latest",
)
p(
    "Endpoint-ul /predict a fost verificat similar, returnând, pe lângă "
    "numărul estimat de zile rămase până la înfundare, coeficientul R² "
    "al regresiei aplicate pe istoricul de citiri acumulat. Răspunsul "
    "real, obținut într-o rulare de testare, este redat mai jos."
)
code(
    """{"status":"ok","days_remaining":2.35,"current_pressure_bar":0.348,
 "clog_threshold_bar":1.5,"degradation_rate_per_hour":0.025856,
 "r_squared":0.0126,"points_used":2077}""",
    label="Raspuns real al endpoint-ului /predict",
)
figure(
    os.path.join(HERE, "screenshot_predict_api.jpg"),
    "Răspunsul JSON real al endpoint-ului /predict, afișat direct în "
    "browser",
    width_in=6.3,
)
p(
    "Valoarea neobișnuit de mică a coeficientului R² (0,0126) din acest "
    "răspuns are o explicație directă, nu reprezintă o eroare a "
    "modelului: fereastra implicită de 24 de ore folosită de /predict "
    "acoperea, la momentul acestei interogări, atât porțiunea "
    "crescătoare a presiunii, cât și repornirea bruscă a senzorului "
    "(vizibilă și în Figura 3.1 și Figura 3.3) — regresia liniară "
    "aplicată pe logaritmul presiunii presupune o singură tendință "
    "exponențială continuă, ipoteză încălcată atunci când datele conțin "
    "un salt discontinuu. Situația confirmă, practic, limitarea "
    "discutată teoretic în secțiunea 5.3: modelul curent nu detectează "
    "automat o schimbare de regim în evoluția presiunii."
)
p(
    "Momentul exact al reset-ului este surprins direct în jurnalul de "
    "execuție al senzorului virtual (Figura 4.2): presiunea diferențială "
    "urcase până la 1,33–1,4 bar (aproape de pragul de înfundare de 1,5 "
    "bar, cu debitul scăzut corespunzător la 6,1–6,7 L/min și turbiditatea "
    "crescută la 3,1–3,3 NTU), moment în care senzorul a fost oprit "
    "manual („[*] Oprire senzor virtual.”) și repornit — repornire "
    "vizibilă prin revenirea presiunii la valoarea de bază, 0,19–0,22 "
    "bar, și afișarea mesajului „[*] Senzor virtual pornit.”"
)
figure(
    os.path.join(HERE, "screenshot_terminal_sensor_reset.png"),
    "Jurnalul senzorului virtual, surprinzând momentul opririi (aproape "
    "de pragul de înfundare) și al repornirii",
    width_in=6.3,
)
p(
    "Endpoint-ul /history a fost verificat cu diferite valori ale "
    "parametrului hours (de exemplu, /history?hours=1), confirmându-se "
    "atât numărul corect de citiri returnate — corelat cu intervalul de "
    "timp cerut și frecvența de publicare a senzorului — cât și "
    "structura fiecărei citiri din listă, identică celei returnate de "
    "/latest."
)

h2("4.4. Verificarea dashboard-ului Grafana")
p(
    "Dashboard-ul „Water Filter Monitor” a fost verificat vizual, în "
    "browser, confirmându-se afișarea corectă a celor patru panouri, cu "
    "actualizare la fiecare cinci secunde. Pe durata unei sesiuni de "
    "testare de aproximativ șase ore, presiunea diferențială a crescut "
    "vizibil, de la aproximativ 0,2 bar la peste 0,4 bar, debitul a "
    "scăzut corespunzător, de la aproximativ 13 L/min la 10 L/min, iar "
    "turbiditatea a urmat un trend crescător similar, de la aproximativ "
    "0,8 NTU la 1,4 NTU — evoluție conformă cu modelul matematic de "
    "degradare descris în secțiunea 3.1.1, panoul de stare a filtrului "
    "indicând „OK” pe toată durata intervalului observat, presiunea "
    "nefiind încă suficient de mare pentru a depăși pragul de înfundare."
)
p(
    "Pentru a verifica și comportamentul pe termen scurt al senzorului, "
    "dashboard-ul a fost inspectat suplimentar pe un interval de doar "
    "15 minute (Figura 4.3), în locul intervalului implicit de 6 ore "
    "folosit până acum."
)
figure(
    os.path.join(HERE, "screenshot_grafana_15min.jpg"),
    "Dashboard-ul pe un interval de 15 minute — se distinge clar "
    "zgomotul aleator adăugat fiecărei citiri",
    width_in=6.3,
)
p(
    "La scara de 6 ore, zgomotul aleator introdus de FilterModel (secțiunea "
    "3.1.1) este vizual absorbit de tendința generală, crescătoare, a "
    "presiunii — graficul apare aproape neted. La scara de 15 minute, "
    "însă, fiecare citire individuală devine vizibilă, iar caracterul "
    "aleator al zgomotului (random.uniform, aplicat independent la "
    "fiecare din cele trei mărimi măsurate) este evident: valorile oscilează "
    "de la o citire la alta, fără a urma un pattern regulat, exact "
    "comportamentul așteptat de la o simulare de zgomot de senzor."
)
p(
    "Grafana permite și inspectarea unui singur panou, în mod extins "
    "(opțiunea View din meniul panoului), utilă pentru citirea precisă a "
    "valorilor și a marcajelor temporale individuale (Figura 4.4)."
)
figure(
    os.path.join(HERE, "screenshot_grafana_panel_zoom.jpg"),
    "Panoul „Cădere de presiune” în vizualizare extinsă, cu marcaje "
    "temporale la fiecare minut",
    width_in=6.3,
)
p(
    "Pentru a observa integral tranziția filtrului către starea de "
    "înfundare, factorul time_acceleration a fost mărit temporar de la "
    "500 la 10 000 (sensor/config.yaml), comprimând un ciclu complet de "
    "colmatare de la aproximativ cinci ore reale la aproximativ "
    "cincisprezece minute. Figura 4.5 surprinde exact momentul tranziției: "
    "presiunea diferențială sare de la aproximativ 0,3 bar la peste 15 "
    "bar, debitul se prăbușește practic la zero, iar turbiditatea crește "
    "de peste zece ori — moment în care panoul de stare comută de la "
    "„OK” la „INFUNDAT”, afișat cu fundal roșu de alertă."
)
figure(
    os.path.join(HERE, "screenshot_grafana_clogged.jpg"),
    "Dashboard-ul în momentul înfundării filtrului — toate cele patru "
    "panouri reflectă simultan tranziția",
    width_in=6.3,
)
p(
    "Această tranziție bruscă, mult mai abruptă decât creșterea "
    "graduală observată pe parcursul restului ciclului, este o "
    "consecință directă a modelului exponențial: odată depășit pragul, "
    "presiunea continuă să crească exponențial, fără nicio limită "
    "superioară impusă în model — comportament util pentru a evidenția "
    "vizual momentul înfundării, dar care nu are, dincolo de acest "
    "prag, o interpretare fizică realistă (un filtru real nu ar "
    "continua să acumuleze presiune diferențială la infinit)."
)

h2("4.5. Verificarea integrării cu simularea de rețea")
p(
    "Pentru a confirma că senzorul virtual aplică efectiv delay-urile "
    "extrase din simularea ns-3 (secțiunea 3.7), a fost comparat "
    "intervalul real dintre două publicări succesive cu intervalul "
    "configurat (cinci secunde). Cu senzorul comutat pe scenariul de "
    "rețea congestionată, jurnalul de execuție a confirmat un interval "
    "real de aproximativ 5,04 secunde între publicări:"
)
code(
    """Trimis: {'timestamp': 1786362513.349, ...}
Trimis: {'timestamp': 1786362518.393, ...}   # diferenta: 5.044s
Trimis: {'timestamp': 1786362523.436, ...}   # diferenta: 5.043s""",
    label="Extras din jurnalul senzorului virtual, scenariul congestionat",
)
p(
    "Diferența de aproximativ 40–44 ms față de intervalul configurat "
    "corespunde exact intervalului de latențe generat de scenariul "
    "congestionat (41,2–67,8 ms, cu media de 60,5 ms, conform Tabelului "
    "3.2), confirmând că integrarea dintre simularea de rețea și "
    "senzorul virtual funcționează conform proiectării."
)

h2("4.6. Verificarea sistemului de alertare")
p(
    "Verificarea end-to-end a sistemului de alertare (secțiunea 3.3.6) "
    "presupune, în mod natural, ca presiunea diferențială să atingă "
    "efectiv un prag de alertă — un proces care, chiar și cu factorul de "
    "accelerare a timpului, durează câteva ore. Pentru a verifica rapid "
    "corectitudinea integrării cu cele două servicii externe (Gmail SMTP "
    "și ntfy.sh), fără a aștepta acest interval, modulul a fost testat "
    "izolat, apelat direct în interiorul containerului backend:"
)
code(
    """$ docker exec -it backend python -c \\
    "from alerting import alert_manager; alert_manager._send_alert(80, 1.2, 1.5)\"""",
    label="Test izolat al modulului de alertare",
)
p(
    "Testul a confirmat primirea, în câteva secunde, atât a email-ului "
    "(cu subiectul „[Water Filter Monitor] Filtru de apa la 80% din "
    "capacitate”), cât și a notificării push pe telefon, prin "
    "aplicația ntfy — validând configurarea corectă a credențialelor SMTP "
    "și a topicului ntfy, independent de starea curentă a senzorului "
    "virtual. Ulterior, mecanismul a fost validat și în fluxul real, "
    "prin repornirea backend-ului (care resetează starea internă a "
    "pragurilor deja notificate) cu senzorul virtual deja rulând peste "
    "un prag de alertă — scenariu în care alertele corespunzătoare "
    "tuturor pragurilor deja depășite sunt declanșate imediat la prima "
    "citire nouă primită."
)
figure(
    os.path.join(HERE, "screenshot_alert_email.png"),
    "Email-ul de alertă, primit real în Gmail, la depășirea pragului de 80%",
    width_in=6.3,
)
figure(
    os.path.join(HERE, "screenshot_alert_push.jpeg"),
    "Notificarea push corespunzătoare, primită pe telefon prin aplicația ntfy",
    width_in=3.4,
)
p(
    "Cele două capturi confirmă livrarea, pe ambele canale, a aceluiași "
    "conținut — subiect și corp de mesaj identice — la un interval de "
    "câteva secunde de la depășirea pragului, timp de răspuns adecvat "
    "pentru un scenariu de avertizare în timp aproape real."
)

h2("4.7. Dificultăți întâmpinate și soluții adoptate")
p(
    "Implementarea sistemului a presupus depășirea mai multor dificultăți "
    "tehnice, majoritatea legate de configurarea mediului de execuție "
    "(Docker Desktop, WSL2) și de compatibilitatea versiunilor de "
    "software folosite. Tabelul 4.2 sintetizează problemele întâmpinate "
    "și soluțiile aplicate — documentarea lor are valoare atât practică, "
    "pentru reproducerea configurării de către alți utilizatori, cât și "
    "metodologică, ilustrând procesul real de depanare a unui sistem "
    "distribuit."
)
table_simple(
    ["Problemă întâmpinată", "Soluție adoptată"],
    [
        [
            "Docker Desktop eșua la pornire cu eroarea „WSL2 is unable to "
            "start since virtualization is not enabled”, deși "
            "virtualizarea hardware era activă (confirmat în Task Manager).",
            "Activarea componentelor Windows lipsă (Virtual Machine "
            "Platform) prin comanda wsl --install --no-distribution, "
            "urmată de repornirea sistemului.",
        ],
        [
            "Comanda python nu era recunoscută — Windows redirecționa "
            "către un stub din Microsoft Store, nu către un interpretor "
            "Python real.",
            "Instalarea Python 3.12 prin winget install "
            "Python.Python.3.12.",
        ],
        [
            "Comanda curl din PowerShell returna un avertisment de "
            "securitate, fiind de fapt un alias pentru Invoke-WebRequest.",
            "Folosirea explicită a binarului curl.exe, distinct de "
            "alias-ul PowerShell.",
        ],
        [
            "Dashboard-ul Grafana afișa „No data” pe toate panourile, "
            "după modificarea manuală a identificatorului (uid) "
            "datasource-ului InfluxDB în fișierul de provisionare.",
            "Ștergerea volumului Docker intern al Grafana "
            "(grafana-data) și reprovizionarea completă, curată, din "
            "fișierele de configurare.",
        ],
        [
            "Scriptul de build al ns-3 (versiunea 3.42) eșua cu o eroare "
            "de tip ValueError în modulul argparse, din cauza "
            "incompatibilității cu Python 3.14 (Ubuntu 26.04).",
            "Trecerea la ultima versiune de dezvoltare ns-3 (ramura "
            "master), compatibilă cu versiunile recente de Python.",
        ],
        [
            "Scenariul ns-3 se oprea cu eroarea „Address Collision”, "
            "cauzată de suprapunerea dintre subrețelele alocate "
            "automat de PointToPointStarHelper și subrețeaua hardcodată "
            "a legăturii backhaul.",
            "Alegerea unei subrețele distincte, în afara intervalului "
            "alocat automat, pentru legătura backhaul.",
        ],
        [
            "Traficul de fond constant, folosit inițial pentru a simula "
            "congestia rețelei, producea o coadă stabilă, cu latență "
            "mare dar practic constantă — nu jitter real.",
            "Înlocuirea traficului de fond constant cu trafic în rafale "
            "aleatoare (model OnOff, cu durate exponențiale), care "
            "produce variație reală a latenței.",
        ],
        [
            "La testarea completă a sistemului de alertare, doar alerta "
            "de 80% a ajuns la utilizator; cele de 90% și 100%, "
            "declanșate câteva secunde mai târziu, au eșuat silențios, "
            "cu eroarea „No address associated with hostname” în "
            "log-urile backend-ului — o pană tranzitorie de rezoluție "
            "DNS în interiorul containerului.",
            "Adăugarea unui mecanism de reîncercare automată (3 "
            "încercări, cu pauză între ele) la trimiterea fiecărei "
            "alerte (email și push), astfel încât o întrerupere "
            "tranzitorie de rețea de câteva secunde să nu mai piardă "
            "definitiv o notificare.",
        ],
    ],
    "Dificultati de implementare si solutiile adoptate",
    col_widths_cm=[7.5, 7.0],
)
p(
    "O a doua observație, legată tot de testarea sistemului de alertare, "
    "a apărut la reluarea testului după corectarea problemei de rețea: "
    "notificările primite raportau procente absurde (de exemplu „5548%”), "
    "în loc de valori între 0% și 100%. Cauza a fost identificată rapid: "
    "corpul mesajului afișa raportul brut dintre presiunea curentă și "
    "pragul de înfundare — un raport care, conform modelului exponențial "
    "descris în secțiunea 3.1.1 (fără limită superioară impusă asupra "
    "presiunii după depășirea pragului), poate crește nelimitat. Soluția "
    "a fost separarea celor două informații: subiectul și mesajul "
    "principal raportează întotdeauna pragul rotund, fix, care a "
    "declanșat alerta (80%, 90% sau 100%), în timp ce presiunea reală, "
    "măsurată, este inclusă separat, clar etichetată în bar, ca "
    "informație tehnică suplimentară — o alegere de proiectare care "
    "prioritizează claritatea mesajului pentru utilizatorul final, fără "
    "a ascunde datele brute de la cei interesați de ele."
)
p(
    "Figura 4.8 ilustrează, direct din terminal, rezolvarea celei de-a "
    "patra probleme din Tabelul 4.2 — eroarea de provisionare a "
    "datasource-ului Grafana. Se observă, în partea superioară, jurnalul "
    "de eroare al containerului Grafana la oprire (Datasource "
    "provisioning error: data source not found), urmat de secvența de "
    "comenzi care a rezolvat problema: ștergerea containerului "
    "(docker compose rm -sf grafana), ștergerea volumului intern "
    "(docker volume rm water-filter-monitor_grafana-data) și "
    "repornirea curată a serviciului (docker compose up -d grafana). "
    "Comanda finală, docker compose ps, confirmă reintrarea tuturor "
    "celor patru servicii în starea „Up”."
)
figure(
    os.path.join(HERE, "screenshot_terminal_grafana_fix.png"),
    "Rezolvarea, din terminal, a erorii de provisionare a "
    "datasource-ului Grafana",
    width_in=6.3,
)

h2("4.8. Sinteza rezultatelor experimentale")
p(
    "Rezultatele obținute pe parcursul testării confirmă, în ansamblu, "
    "funcționarea corectă a sistemului proiectat: fluxul de date "
    "circulă fără pierderi de la senzorul virtual până la dashboard, "
    "predicția zilelor rămase până la înfundare reflectă corect "
    "tendința observată în datele istorice, iar integrarea cu "
    "simularea de rețea ns-3 reproduce, la nivel aplicativ, un delay "
    "măsurabil și consistent cu scenariul de rețea simulat. Rezultatele "
    "cantitative obținute din simularea de rețea (secțiunea 3.7.3) "
    "evidențiază o diferență semnificativă, atât ca medie cât și ca "
    "variabilitate, între un scenariu de rețea neîncărcată și unul "
    "congestionat — un rezultat relevant pentru discuția, din capitolele "
    "următoare, privind beneficiile rețelelor de generație nouă (5G) "
    "pentru aplicații IoT sensibile la latență și jitter."
)
p(
    "Capitolul următor al lucrării va prezenta concluziile generale ale "
    "acestui demers, limitările abordării alese și direcțiile posibile "
    "de continuare a lucrării."
)

# ================================================================ CAPITOLUL 5
new_chapter(5)
h1("CAPITOLUL 5. CONCLUZII")

h2("5.1. Concluzii generale")
p(
    "Lucrarea de față a urmărit proiectarea și implementarea unui sistem "
    "IoT complet, de tip edge-to-cloud, pentru monitorizarea stării unui "
    "filtru de apă și predicția momentului de înfundare a acestuia, "
    "însoțit de o evaluare a impactului condițiilor de rețea asupra "
    "livrării datelor, realizată prin simulare cu ns-3. Obiectivele "
    "stabilite în Capitolul 1 au fost îndeplinite: a fost proiectată și "
    "implementată o arhitectură funcțională, structurată pe niveluri "
    "clare (achiziție, transport, procesare, persistență, prezentare), "
    "cu un senzor virtual guvernat de un model matematic de degradare "
    "plauzibil fizic, un model predictiv bazat pe regresie liniară "
    "aplicată logaritmului presiunii diferențiale, un dashboard de "
    "vizualizare în timp real și o simulare de rețea capabilă să "
    "reproducă atât o rețea rapidă, neîncărcată, cât și o rețea "
    "congestionată, cu jitter real."
)
p(
    "Testarea descrisă în Capitolul 4 a confirmat funcționarea corectă a "
    "întregului lanț de procesare a datelor, de la generarea unei citiri "
    "de către senzorul virtual, până la afișarea ei pe dashboard și "
    "calcularea unei predicții asupra stării filtrului. Rezultatele "
    "cantitative obținute din simularea de rețea au evidențiat o "
    "diferență semnificativă — atât ca latență medie, cât și ca "
    "variabilitate (jitter) — între un scenariu de rețea rapidă și unul "
    "congestionat, confirmând relevanța practică a temei alese și "
    "oferind un argument concret în favoarea beneficiilor rețelelor de "
    "generație nouă pentru aplicații IoT sensibile la promptitudinea "
    "livrării datelor."
)

h2("5.2. Contribuții și rezultate obținute")
p(
    "Principalele rezultate concrete obținute în cadrul acestei lucrări "
    "sunt sintetizate mai jos:"
)
bullets(
    [
        "O arhitectură IoT edge-to-cloud complet funcțională, "
        "containerizată și orchestrată printr-o singură comandă "
        "(Docker Compose), reproductibilă pe orice sistem de calcul "
        "compatibil.",
        "Un model matematic de degradare a filtrului, cu bază fizică "
        "plauzibilă (creștere exponențială a presiunii diferențiale), "
        "implementat într-un senzor virtual configurabil.",
        "Un model de predicție bazat pe regresie liniară, care "
        "estimează numărul de zile rămase până la înfundarea filtrului "
        "direct din istoricul de citiri, cu raportarea calității "
        "potrivirii (R²).",
        "Un dashboard de vizualizare în timp real, provizionat complet "
        "automat, fără configurare manuală.",
        "Un scenariu de simulare de rețea realizat cu ns-3, capabil să "
        "reproducă atât o rețea rapidă, neîncărcată, cât și o rețea "
        "congestionată, cu jitter real, generat prin trafic de fond în "
        "rafale aleatoare — o metodologie mai riguroasă decât o simplă "
        "întârziere fixă, adesea folosită în abordări simplificate.",
        "Un sistem de alertare pe două canale (email și notificare push "
        "pe telefon, prin ntfy.sh), care avertizează activ utilizatorul "
        "la trecerea unor praguri de colmatare (80%, 90%, înfundare "
        "completă), cu deduplicare a alertelor și resetare automată la "
        "înlocuirea filtrului.",
        "Documentație tehnică completă a arhitecturii, codului sursă și "
        "procesului de testare, publicată public, însoțită de un "
        "tutorial de instalare și utilizare pas cu pas.",
    ]
)

h2("5.3. Limitările abordării")
p(
    "Deși obiectivele propuse au fost atinse, abordarea aleasă prezintă "
    "o serie de limitări, asumate conștient încă din faza de proiectare "
    "și discutate, parțial, în secțiunea 1.3:"
)
bullets(
    [
        "Senzorul virtual simulează degradarea filtrului printr-un "
        "model matematic unic (creștere exponențială a presiunii), care "
        "nu acoperă alte moduri posibile de degradare a unui filtru "
        "real (de exemplu, colmatare neuniformă, deteriorarea "
        "materialului filtrant, variații ale calității apei de la "
        "sursă) — validarea modelului împotriva unor date reale, de la "
        "un filtru fizic instrumentat, rămâne un pas necesar pentru "
        "aplicabilitatea practică a sistemului.",
        "Modelul de predicție presupune că tendința exponențială "
        "observată se menține până la atingerea pragului de înfundare; "
        "în situații reale, rata de degradare se poate schimba brusc "
        "(de exemplu, în urma unei schimbări a calității apei de la "
        "sursă), situație pe care regresia liniară curentă nu o "
        "detectează automat.",
        "Simularea de rețea realizată cu ns-3 nu este integrată live cu "
        "restul sistemului, ci rulează separat, iar rezultatele sunt "
        "reproduse ca fișiere statice — o decizie asumată pentru "
        "simplitate și robustețe, dar care nu reflectă variabilitatea "
        "unei rețele reale, în timp real, condiționată de trafic extern "
        "sistemului monitorizat.",
        "Scenariul de rețea simulat folosește module ns-3 generice "
        "(point-to-point, cu parametri configurați pentru a aproxima "
        "caracteristici tipice 5G), nu modulul dedicat de simulare 5G "
        "NR (nr / 5G-LENA) — o alegere motivată de complexitatea "
        "semnificativ mai mare de instalare și configurare a acestuia "
        "din urmă, discutată în etapa de planificare a acestei "
        "componente.",
        "Configurația actuală a componentelor de infrastructură "
        "(autentificare anonimă la broker-ul MQTT, credențiale "
        "implicite pentru InfluxDB și Grafana) este adecvată unui mediu "
        "de dezvoltare și demonstrație, dar nu unui mediu de producție, "
        "unde ar fi necesare măsuri suplimentare de securitate "
        "(autentificare, criptare TLS, rotația token-urilor de acces).",
        "Testarea a fost realizată la scară redusă (un singur senzor "
        "virtual, respectiv șase noduri în simularea de rețea) — "
        "comportamentul sistemului la o scară mult mai mare (sute sau "
        "mii de senzori) nu a fost evaluat.",
    ]
)

h2("5.4. Direcții de dezvoltare ulterioară")
p(
    "Pe baza limitărilor identificate, pot fi conturate mai multe "
    "direcții de continuare a acestei lucrări:"
)
bullets(
    [
        "Integrarea unui senzor fizic real (de exemplu, un senzor de "
        "presiune diferențială montat pe un filtru de apă instrumentat), "
        "care să publice date în același format, pe același topic MQTT, "
        "fără nicio modificare a restului sistemului — validând, în "
        "condiții reale, atât modelul de degradare, cât și cel de "
        "predicție.",
        "Extinderea modelului de predicție cu tehnici capabile să "
        "surprindă schimbări de regim ale ratei de degradare (de "
        "exemplu, regresie polinomială pe ferestre glisante, sau modele "
        "de tip rețea neuronală recurentă, pentru scenarii de "
        "degradare neliniare).",
        "Realizarea unei simulări de rețea mai riguroase, folosind "
        "modulul dedicat 5G NR pentru ns-3 (nr / 5G-LENA), sau "
        "validarea rezultatelor simulate împotriva unor măsurători "
        "reale, efectuate pe o rețea 5G comercială.",
        "Extinderea sistemului de alertare (implementat, secțiunea "
        "3.3.6) cu un canal de tip reminder programat în avans — de "
        "exemplu, crearea automată a unui eveniment în Google Calendar, "
        "la data estimată de predicție, pentru a oferi utilizatorului un "
        "orizont de planificare, nu doar o alertă reactivă la depășirea "
        "unui prag.",
        "Extinderea arhitecturii pentru gestionarea unei flote de "
        "filtre/senzori multipli, cu un dashboard agregat la nivel de "
        "flotă, și evaluarea comportamentului sistemului la scară "
        "mărită.",
        "Consolidarea securității componentelor de infrastructură "
        "(autentificare MQTT, criptare TLS, gestionarea centralizată a "
        "secretelor), ca pas necesar către un mediu de producție.",
    ]
)

h2("5.5. Cuvinte de închidere")
p(
    "Sistemul realizat în cadrul acestei lucrări demonstrează, într-un "
    "cadru practic și reproductibil, modul în care mai multe direcții "
    "tehnologice actuale — comunicația IoT de tip publish/subscribe, "
    "stocarea și interogarea seriilor de timp la scară, predicția bazată "
    "pe modele statistice interpretabile, vizualizarea datelor în timp "
    "real și simularea rețelelor de nouă generație — pot fi integrate "
    "coerent într-un singur studiu de caz, aplicabil unei probleme "
    "concrete și relevante: mentenanța predictivă a unui echipament "
    "casnic sau industrial. Limitările asumate și direcțiile de "
    "dezvoltare identificate în acest capitol conturează, totodată, un "
    "traseu clar de continuare a lucrării, către o soluție validată în "
    "condiții reale de utilizare."
)

# ================================================================ BIBLIOGRAFIE
para = doc.add_paragraph()
para.paragraph_format.page_break_before = True
bib = doc.add_paragraph()
bib.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = bib.add_run("BIBLIOGRAFIE")
set_run_font(r, size=18, bold=True)
bib.paragraph_format.space_after = 20

p(
    "Documentația oficială a tehnologiilor folosite în implementarea "
    "sistemului, consultată pe parcursul acestei lucrări:"
)
bib_items = [
    "FastAPI — https://fastapi.tiangolo.com/",
    "InfluxDB 2.x / limbajul Flux — https://docs.influxdata.com/",
    "Grafana — https://grafana.com/docs/grafana/latest/",
    "Eclipse Mosquitto / protocolul MQTT — https://mqtt.org/ , "
    "https://mosquitto.org/documentation/",
    "ns-3 Network Simulator — https://www.nsnam.org/documentation/",
    "Docker / Docker Compose — https://docs.docker.com/",
    "scikit-learn — https://scikit-learn.org/stable/",
    "Windows Subsystem for Linux (WSL2) — https://learn.microsoft.com/windows/wsl/",
]
bullets(bib_items)

p(
    "[Secțiune de completat de autor cu resursele bibliografice "
    "suplimentare — cărți, articole și lucrări științifice — consultate "
    "pentru fundamentarea teoretică a temei: IoT și mentenanță "
    "predictivă, protocoale de comunicație pentru IoT, baze de date "
    "de tip serie-de-timp, tehnici de regresie statistică, rețele "
    "mobile 5G.]",
    italic=True,
)

doc.save(os.path.join(HERE, "Lucrare_Disertatie.docx"))
print("Document salvat cu succes.")
